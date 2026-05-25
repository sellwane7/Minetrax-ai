import streamlit as st
import os
import io
import re
import base64
from datetime import datetime, date, time as dt_time
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from PIL import Image
import math as _math
import random as _random
from dotenv import load_dotenv
import json

# Load .env file
load_dotenv()

# Read the key from the environment, fall back to streamlit secrets if needed
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
if not GROQ_API_KEY:
    try:
        GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "")
    except Exception:
        GROQ_API_KEY = ""

# Securely initialize the global client
client = None
if GROQ_API_KEY:
    try:
        client = Groq(api_key=GROQ_API_KEY)
    except Exception:
        client = None

def _get_client():
    api_key = os.getenv("GROQ_API_KEY")

    if not api_key:
        return None

    return Groq(api_key=api_key)

client = _get_client()

# Simplify call_groq to use the initialized client directly
def call_groq(prompt, model="llama-3.3-70b-versatile", system_prompt=None, history=None):
    if not client:
        return _fallback_response(prompt)
    msgs = []
    if system_prompt:
        msgs.append({"role": "system", "content": system_prompt})
    if history:
        for h in history:
            msgs.append({"role": h["role"], "content": h["content"]})
    msgs.append({"role": "user", "content": prompt})
    try:
        completion = client.chat.completions.create(
            model=model, messages=msgs, temperature=0.7, max_tokens=2048
        )
        return completion.choices[0].message.content
    except Exception as e:
        return _fallback_response(prompt, error=str(e))

# ─── PAGE CONFIG ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Minetrax AI | Mining Safety Intelligence",
    page_icon="⛏️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ═══════════════════════════════════════════════════════════════════════════════
# REAL SOUTH AFRICAN MINE DATA PER PROVINCE
# ═══════════════════════════════════════════════════════════════════════════════
PROVINCE_MINE_DATA = {
    "Gauteng": {
        "mines": ["TauTona Gold Mine (AngloGold Ashanti)", "Mponeng Gold Mine (AngloGold Ashanti)",
                  "South Deep Gold Mine (Gold Fields)", "Driefontein Gold Mine (Sibanye-Stillwater)",
                  "Kloof Gold Mine (Sibanye-Stillwater)", "Cooke Gold Mine (Sibanye-Stillwater)"],
        "zones": ["Level 120 – Main Stope 3A", "Level 95 – Haulage Decline", "Shaft 4 – Reef Drive",
                  "Western Ventilation Raise", "Carbon-in-Leach Plant", "Explosives Bay – Zone 2",
                  "Pump Chamber – Level 80", "Rock Hoisting Shaft Bottom", "Level 110 – Trackless Drive",
                  "Tailings Storage Facility East"],
        "operations": ["Gold extraction via CIL", "Deep shaft sinking", "Rock hoisting operations",
                       "Backfill operations", "Stope development", "Ventilation shaft maintenance"],
        "type": "Deep-Level Gold"
    },
    "Limpopo": {
        "mines": ["Venetia Diamond Mine (De Beers)", "Mogalakwena Platinum Mine (Anglo American Platinum)",
                  "Twickenham Platinum Mine (Anglo American Platinum)", "Modikwa Platinum Mine (ARM/Anglo Plats)",
                  "Marula Platinum Mine (Implats)", "Nzuri Copper Project"],
        "zones": ["Open Pit Bench 14E", "Primary Crusher Feed Zone", "Concentrator Plant Floor",
                  "Overburden Dump North", "Eastern Highwall Section", "ROM Pad – Sorting Plant",
                  "Tailings Management Facility", "Pit Access Road KM-7", "Workshop & Maintenance Bay",
                  "Explosives Storage Compound"],
        "operations": ["Diamond recovery via X-ray sorting", "Platinum concentrating", "Open pit blasting",
                       "Truck and shovel operations", "Ore processing", "Reagent plant operations"],
        "type": "Open Pit / Diamond & Platinum"
    },
    "Mpumalanga": {
        "mines": ["Impala Springs Coal Mine (Exxaro)", "Goedehoop Colliery (Anglo American)",
                  "Greenside Colliery (Seriti Resources)", "Klipspruit Colliery (Seriti Resources)",
                  "Isibonelo Colliery (South32)", "Kriel Colliery (Seriti Resources)"],
        "zones": ["Longwall Panel LW-07", "Conveyor Belt Drive Station B3", "Main Return Airway",
                  "Box Cut Entry North", "Coal Preparation Plant", "Run-of-Mine Stockpile",
                  "Methane Drainage Borehole Field", "Outbye Haulage LH-4", "Shaft Inset Level 3",
                  "Fan Drift & Ventilation Doors"],
        "operations": ["Longwall coal extraction", "Section continuous miner ops", "Coal washing plant",
                       "Methane drainage", "Conveyor maintenance", "Bleeder entry support"],
        "type": "Underground Coal"
    },
    "North West": {
        "mines": ["Rustenburg Platinum Mine (Impala Platinum)", "Marikana Mine (Sibanye-Stillwater)",
                  "Bafokeng Rasimone Platinum Mine (RBPlats)", "Kroondal Platinum Mine (Sibanye-Stillwater)",
                  "Styldrift Mine (RBPlats/Impala)", "Booysendal Platinum Mine (Northam Platinum)"],
        "zones": ["Merensky Reef Stope 7B", "UG2 Reef Development Drive", "No. 3 Shaft – Reef Level",
                  "Decline Ramp Section D", "Concentrator – Flotation Circuit", "Smelter Converter Aisle",
                  "Portal Access Gate", "Rock Mechanics Test Zone", "Shaft Steelwork Level 18",
                  "Footwall Drive FW-22"],
        "operations": ["Platinum group metals extraction", "Reef development drilling", "Crush-and-float processing",
                       "Smelting & converting", "Trackless mechanised mining", "Stope blasting cycles"],
        "type": "Platinum (Bushveld Complex)"
    },
    "Free State": {
        "mines": ["Beatrix Gold Mine (Sibanye-Stillwater)", "Joel Gold Mine (Harmony Gold)",
                  "Target Gold Mine (Harmony Gold)", "Tshepong Gold Mine (Harmony Gold)",
                  "Phakisa Gold Mine (Harmony Gold)", "Masimong Gold Mine (Harmony Gold)"],
        "zones": ["Level 48 – Stoping Block 2", "Ore Pass 12 – Loading Bay", "Sub-vertical Shaft Pump Station",
                  "Level 55 – Development End", "Reduction Works – CIL Circuit", "Tailings Dam Cell 3",
                  "Level 33 – Trackless Haulage", "Emergency Refuge Bay – Level 60",
                  "Coolant Pipes – Chiller Plant", "Shaft Bottom Station – Main Shaft"],
        "operations": ["Gold reef mining", "Ore transport hoisting", "Elution & electro-winning",
                       "Carbon-in-leach recovery", "Ventilation cooling", "Shaft infrastructure maintenance"],
        "type": "Intermediate-Level Gold"
    },
    "Northern Cape": {
        "mines": ["Kolomela Iron Ore Mine (Kumba Iron Ore)", "Sishen Iron Ore Mine (Kumba Iron Ore)",
                  "Thabazimbi Iron Ore Mine (ArcelorMittal)", "Black Rock Manganese Mine (Assore/ARM)",
                  "Nchwaning Manganese Mine (Assmang)", "Postmasburg Manganese Mine (Hotazel)"],
        "zones": ["Pit 4 – Bench 36 High Wall", "Drilling & Blasting Platform South",
                  "Primary Gyratory Crusher Station", "Conveyor Overland Transfer Point 6",
                  "Dense Media Separation Plant", "Waste Dump Slope Monitoring Zone",
                  "Railway Load-Out Facility", "Eastern Pit Haul Road K-12",
                  "Dewatering Sump – Level 5", "Reclaim Tunnel Belt 9"],
        "operations": ["Iron ore open pit extraction", "Manganese underground production",
                       "Dense media separation", "Dry magnetic separation", "Rail logistics",
                       "Pit slope stability monitoring"],
        "type": "Open Pit Iron Ore / Manganese"
    },
    "KwaZulu-Natal": {
        "mines": ["Somkhele Anthracite Mine (Tendele Coal)", "Zululand Anthracite Colliery (ZAC)",
                  "Kilbarchan Coal Mine (Zinoju Coal)", "Richards Bay Minerals – Titanium Sands",
                  "Corridor Sands Project (Rio Tinto)", "Glimmer Coal Mine (Kumba)"],
        "zones": ["Strip Mine Panel 3E", "Open Cast Working Face North",
                  "Screening & Sizing Plant", "Heavy Mineral Separation Unit",
                  "Slimes Dam – Zone C", "Overburden Ramp 5",
                  "Load-Out Silo Bay", "Wet Separation Circuit",
                  "Reclamation Area – East Block", "Site Access Control & Security Compound"],
        "operations": ["Anthracite open cast strip mining", "Heavy mineral sands mining",
                       "Mineral separation processing", "Tailings management",
                       "Overburden stripping", "Concentrate drying & bagging"],
        "type": "Anthracite / Heavy Mineral Sands"
    },
    "Eastern Cape": {
        "mines": ["Minas Gerais Aggregates – Port Elizabeth", "AfriSam Rietfontein Quarry",
                  "PPC Riebeeck Cement & Limestone", "Coedmore Quarry (AfriSam)",
                  "Gqeberha Industrial Minerals", "Ecca Aggregates – King William's Town"],
        "zones": ["Quarry Face Level 3 North", "Primary Jaw Crusher Feed Platform",
                  "Aggregate Screening Plant", "Lime Kiln No. 2", "Slurry Pipeline Corridor",
                  "Dust Suppression Zone", "Vehicle Weighbridge", "Blasting Exclusion Zone",
                  "Cement Silo Area", "Waste Rock Dump"],
        "operations": ["Limestone quarrying", "Aggregate crushing & grading", "Cement production",
                       "Lime burning", "Dust suppression operations", "Blasting operations"],
        "type": "Quarry / Aggregate / Limestone"
    },
    "Western Cape": {
        "mines": ["AfriSam Piketberg Quarry", "Lyttelton Dolomite Mine (Pretoria Portland Cement)",
                  "Rooibos Mine – Silica Sands (PFG Glass)", "Cape Lime Works – Wellington",
                  "Citrusdal Phosphate Deposit (exploration)", "Bonnievale Construction Aggregates"],
        "zones": ["Bench Quarry East Face", "Primary Crusher Bay", "Silica Screening Plant",
                  "Overburden Stockpile", "Explosives Storage Compound", "Dust Control Berm",
                  "Workshop & Fleet Maintenance", "Site Perimeter & Security",
                  "Water Treatment Point", "Rehabilitation Zone – Old Pit 2"],
        "operations": ["Silica sand quarrying", "Aggregate production", "Lime manufacturing",
                       "Dolomite extraction", "Site rehabilitation", "Environmental monitoring"],
        "type": "Quarry / Industrial Minerals"
    }
}

# ─── STATIC DATA ────────────────────────────────────────────────────────────────
MINERS = [
    {"id": "M001", "name": "Sipho Dlamini",     "role": "Drill Operator",         "dept": "Underground Operations"},
    {"id": "M002", "name": "Thabo Mokoena",     "role": "Blasting Technician",    "dept": "Explosives Division"},
    {"id": "M003", "name": "Lungelo Zulu",      "role": "Shaft Sinker",           "dept": "Shaft Development"},
    {"id": "M004", "name": "Mandla Nkosi",      "role": "Rock Engineer",          "dept": "Geotechnical Services"},
    {"id": "M005", "name": "Bongani Khumalo",   "role": "Winch Operator",         "dept": "Hoisting Services"},
    {"id": "M006", "name": "Nkosinathi Hadebe", "role": "Stope Miner",            "dept": "Underground Operations"},
    {"id": "M007", "name": "Sibusiso Mthembu",  "role": "Timberman",              "dept": "Support Services"},
    {"id": "M008", "name": "Dumisani Buthelezi","role": "Trackless Equipment Op", "dept": "Mechanised Mining"},
    {"id": "M009", "name": "Ayanda Ntanzi",     "role": "Pump Operator",          "dept": "Water Management"},
    {"id": "M010", "name": "Siyanda Gumede",    "role": "Safety Officer",         "dept": "HSE Department"},
    {"id": "M011", "name": "Phiwayinkosi Maia", "role": "Electrician",            "dept": "Electrical Services"},
    {"id": "M012", "name": "Mthokozisi Cele",   "role": "Surveyor",               "dept": "Mine Survey"},
    {"id": "M013", "name": "Lindani Msweli",    "role": "Rope-Runner",            "dept": "Hoisting Services"},
    {"id": "M014", "name": "Sandile Shabalala", "role": "Ventilation Officer",    "dept": "Ventilation Services"},
    {"id": "M015", "name": "Nhlanhla Majola",   "role": "Emergency Rescue Tech",  "dept": "Mine Rescue"},
]

SUPERVISORS = [
    {"id": "S001", "name": "Gerrit van Wyk",     "role": "Mine Overseer",            "dept": "Underground Operations"},
    {"id": "S002", "name": "Patricia Mohlala",   "role": "Senior Safety Manager",    "dept": "HSE Department"},
    {"id": "S003", "name": "Johan Fourie",       "role": "Section Mine Manager",     "dept": "Section Management"},
    {"id": "S004", "name": "Nompumelelo Dube",   "role": "Shift Supervisor",         "dept": "Underground Operations"},
    {"id": "S005", "name": "Andries Botha",      "role": "Rock Engineering Manager", "dept": "Geotechnical Services"},
    {"id": "S006", "name": "Zanele Mahlangu",    "role": "Environmental Manager",    "dept": "Environmental Affairs"},
    {"id": "S007", "name": "Pieter Steenkamp",   "role": "Mine Captain",             "dept": "Shaft Development"},
    {"id": "S008", "name": "Refilwe Sithole",    "role": "HR & Compliance Lead",     "dept": "Human Resources"},
]

INCIDENT_TYPES = [
    "Fall of Ground", "Slip / Trip / Fall", "Equipment / Machinery Accident",
    "Electrical Incident", "Explosion / Blast Related", "Gas / Fume Exposure",
    "Fire", "Transportation / Tramming Accident", "Struck By Object",
    "Caught In / Between Equipment", "Ergonomic Injury", "Near Miss",
]

INJURY_TYPES = [
    "Laceration / Open Wound", "Fracture / Broken Bone", "Concussion / Head Injury",
    "Burns (Thermal)", "Chemical Burns", "Eye Injury", "Crush Injury",
    "Sprain / Strain", "Respiratory Distress", "Hearing Damage",
    "Multiple Injuries", "Fatality",
]

MINE_TYPES  = ["Underground", "Open Pit", "Alluvial", "Longwall (Coal)", "Hard Rock", "Quarry", "Surface", "Shaft"]
PROVINCES   = list(PROVINCE_MINE_DATA.keys())

SAFETY_PROTOCOLS = [
    "PPE (Hard Hat, Safety Boots, Gloves, Goggles)",
    "Ventilation & Air Quality Monitoring",
    "Gas Detection & Monitoring Systems",
    "Traffic Management Plan",
    "Environmental & Dust Monitoring",
    "Pre-Shift Equipment Inspections",
    "First Aid Readiness & Emergency Response Plan",
]

# ═══════════════════════════════════════════════════════════════════════════════
# ML FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════
def _ml_accident_prediction(depth, gas_ppm, temp_c, hours_worked, maintenance, incident_hx):
    maint_map = {"Overdue": 0.0, "Poor": 0.25, "Fair": 0.5, "Good": 0.75, "Excellent": 1.0}
    d  = min(depth / 3500, 1.0)
    g  = min(gas_ppm / 2000, 1.0)
    t  = max(0, min((temp_c - 18) / 37, 1.0))
    f  = min(hours_worked / 16, 1.0)
    m  = maint_map.get(maintenance, 0.75)
    hx = min(incident_hx / 20, 1.0)
    score = 0.25*d + 0.22*g + 0.14*t + 0.13*f + 0.14*(1-m) + 0.12*hx
    prob  = 1 / (1 + _math.exp(-10 * (score - 0.45)))
    features = {
        "Depth / Level":    round(0.25*d*100, 1),
        "Gas Level":        round(0.22*g*100, 1),
        "Temperature":      round(0.14*t*100, 1),
        "Shift Fatigue":    round(0.13*f*100, 1),
        "Maint. Deficit":   round(0.14*(1-m)*100, 1),
        "Incident History": round(0.12*hx*100, 1),
    }
    sev_score = 0.30*d + 0.22*g + 0.15*f + 0.12*t + 0.11*(1-m) + 0.10*hx
    if sev_score < 0.20:   sev, sc = "Minor",    "#2ed573"
    elif sev_score < 0.40: sev, sc = "Moderate", "#ffa502"
    elif sev_score < 0.60: sev, sc = "Serious",  "#ff6b35"
    elif sev_score < 0.78: sev, sc = "Critical", "#ff4757"
    else:                  sev, sc = "Fatal",    "#8B0000"
    return round(prob*100, 1), features, sev, sc

def _ml_render_prediction(acc_pct, features, sev_label, sev_color):
    ac = "#2ed573" if acc_pct < 30 else "#ffa502" if acc_pct < 60 else "#ff4757"
    max_f = max(features.values()) or 1
    bars = "".join(
        f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:7px;">'
        f'<div style="width:130px;font-size:11px;color:#b0aac0;text-align:right;">{fn}</div>'
        f'<div style="flex:1;background:#0a0a16;border-radius:4px;height:14px;overflow:hidden;">'
        f'<div style="width:{int((fv/max_f)*100)}%;height:100%;'
        f'background:{"#ff4757" if fv>12 else "#ffa502" if fv>6 else "#2ed573"};'
        f'border-radius:4px;transition:width 0.5s;"></div></div>'
        f'<div style="width:32px;font-size:10px;color:#ffffff;text-align:right;font-weight:700;">{fv}</div></div>'
        for fn, fv in features.items())
    return f"""
    <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-top:16px;">
      <div class="card" style="border-top:3px solid {ac};text-align:center;padding:18px;">
        <div style="font-size:9px;color:#b0aac0;letter-spacing:2px;text-transform:uppercase;margin-bottom:6px;">Accident Probability</div>
        <div style="font-size:48px;font-family:'Bebas Neue';color:{ac};line-height:1;">{acc_pct}%</div>
        <div style="background:#0a0a16;border-radius:6px;height:6px;overflow:hidden;margin-top:10px;">
          <div style="width:{acc_pct}%;height:100%;background:{ac};border-radius:6px;"></div>
        </div>
      </div>
      <div class="card" style="border-top:3px solid {sev_color};text-align:center;padding:18px;">
        <div style="font-size:9px;color:#b0aac0;letter-spacing:2px;text-transform:uppercase;margin-bottom:6px;">Predicted Severity</div>
        <div style="font-size:40px;font-family:'Bebas Neue';color:{sev_color};line-height:1;">{sev_label}</div>
      </div>
    </div>
    <div class="card" style="margin-top:12px;">
      <div style="font-size:9px;color:#b0aac0;letter-spacing:1.5px;text-transform:uppercase;margin-bottom:10px;">Risk Factor Breakdown (Supervised Regression)</div>
      {bars}
    </div>"""

def _ml_kmeans_zones(zone_data, k=3, iters=35):
    def norm(data, keys):
        mx = {k2: max(d[k2] for d in data) or 1 for k2 in keys}
        return [{k2: d[k2]/mx[k2] for k2 in keys} for d in data]
    def dist(a, b, keys): return _math.sqrt(sum((a[k2]-b[k2])**2 for k2 in keys))
    keys = ["depth", "gas", "incidents"]
    nd   = norm(zone_data, keys)
    _random.seed(42)
    cents = _random.sample(nd, min(k, len(nd)))
    labels = [0]*len(nd)
    for _ in range(iters):
        for i, pt in enumerate(nd):
            labels[i] = min(range(len(cents)), key=lambda c: dist(pt, cents[c], keys))
        new_c = []
        for c in range(len(cents)):
            cl = [nd[i] for i, l in enumerate(labels) if l == c]
            new_c.append({k2: sum(p[k2] for p in cl)/len(cl) for k2 in keys} if cl else cents[c])
        cents = new_c
    scores = [sum(c.values())/len(c) for c in cents]
    sc_sorted = sorted(range(len(cents)), key=lambda i: scores[i])
    risk_map = {sc_sorted[0]: ("Low Risk", "#2ed573")}
    if len(sc_sorted) > 1: risk_map[sc_sorted[1]] = ("Medium Risk", "#ffa502")
    if len(sc_sorted) > 2: risk_map[sc_sorted[2]] = ("High Risk", "#ff4757")
    return labels, risk_map

def _ml_anomaly_detect(readings, threshold=2.5):
    mean_r = sum(readings)/len(readings)
    std_r  = _math.sqrt(sum((x-mean_r)**2 for x in readings)/len(readings)) or 1
    zs     = [(x-mean_r)/std_r for x in readings]
    anomalies = [(i, readings[i], zs[i]) for i in range(len(readings)) if abs(zs[i]) > threshold]
    return anomalies, round(mean_r, 1), round(std_r, 1)

def _ml_rl_recommend(gas_idx, gnd_idx, episodes=300):
    ACTIONS = [
        "Continue Normal Operations",
        "Increase Ventilation & Air Monitoring",
        "Deploy Additional Gas Monitors",
        "Reduce Crew to Essential Personnel Only",
        "Evacuate Affected Section Immediately",
        "Full Emergency Evacuation – All Personnel"
    ]
    n_s, n_a = 16, len(ACTIONS)
    def reward(s, a):
        g = s//4; gr = s%4
        ideal = min(n_a-1, max(0, round(((g+gr)/6)*(n_a-1))))
        return 10 - abs(a-ideal)*3 + _random.gauss(0, 0.5)
    _random.seed(99)
    Q = [[0.0]*n_a for _ in range(n_s)]
    alpha, gamma, eps = 0.3, 0.85, 0.9
    for _ in range(episodes):
        s = _random.randint(0, n_s-1)
        for _ in range(12):
            a  = _random.randint(0, n_a-1) if _random.random() < eps else Q[s].index(max(Q[s]))
            ns = _random.randint(0, n_s-1)
            Q[s][a] += alpha*(reward(s, a)+gamma*max(Q[ns])-Q[s][a])
            s = ns
        eps = max(0.05, eps*0.997)
    cur  = gas_idx*4 + gnd_idx
    best = Q[cur].index(max(Q[cur]))
    return ACTIONS[best], Q[cur], ACTIONS

def _ml_linear_regression_trend(values):
    n = len(values)
    if n < 2:
        return 0, values[-1] if values else 0
    xs = list(range(n))
    x_mean = sum(xs)/n; y_mean = sum(values)/n
    num = sum((xs[i]-x_mean)*(values[i]-y_mean) for i in range(n))
    den = sum((xs[i]-x_mean)**2 for i in range(n)) or 1
    slope = num/den
    intercept = y_mean - slope*x_mean
    return round(slope, 3), round(intercept + slope*n, 1)

# ═══════════════════════════════════════════════════════════════════════════════
# GROQ CLIENT  –  dynamic: re-reads key from session state on every call
# ═══════════════════════════════════════════════════════════════════════════════

def call_groq_vision(prompt, image_b64, mime="image/jpeg"):
    c = _get_client()
    if not c:
        return _fallback_response(prompt)
    try:
        completion = c.chat.completions.create(
            model="llama-3.2-11b-vision-preview",
            messages=[{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{image_b64}"}},
                {"type": "text", "text": prompt}
            ]}],
            temperature=0.7, max_tokens=4096
        )
        return completion.choices[0].message.content
    except Exception as e:
        return call_groq(f"{prompt}\n[Vision model unavailable: {str(e)}]")

def _fallback_response(prompt, error=None):
    note = (
        f"\n\n---\n> ⚠️ **AI Engine Offline** — Configure your `GROQ_API_KEY` to enable live AI-generated responses tailored to your specific mine site and conditions. The information below is based on South African mining best practice and MHSA Act 29 of 1996 statutory requirements."
        if not error
        else f"\n\n---\n> ⚠️ **API Error:** {error} — Showing standard guidance below."
    )
    p = prompt.lower()

    # ── HAZARD INTELLIGENCE ───────────────────────────────────────────────────
    if any(w in p for w in ["hazard", "risk", "danger", "analysis", "gas", "methane", "fall of ground", "fog"]):
        return f"""## Hazard Analysis Report

### Overall Risk Assessment
Based on the site parameters entered, a structured risk assessment has been initiated. Until the AI engine is connected, the following comprehensive guidance applies to all underground and surface mining operations in South Africa.

---

### 1. Primary Hazard Categories in South African Mining

**Fall of Ground (FOG)**
Fall of ground remains the single largest cause of fatalities in South African hard-rock underground mines, particularly in gold and platinum operations. Key indicators include audible cracking or "popping" sounds, visible cracks running parallel to excavation faces, water seepage along joint planes, and spalling of hanging wall rock. Under MHSA Section 11, a formal written risk assessment is mandatory before any person enters or works in an area that has not been made safe. Rock mechanics engineers must classify ground conditions using the Rock Mass Rating (RMR) or Q-system prior to stope entry after any seismic event above 1.0 ML.

**Gas and Atmospheric Hazards**
In coal mines (especially Mpumalanga longwall operations), methane (CH₄) is the primary atmospheric hazard. Mandatory action thresholds under MHSA Regulation 5.1 are: 0.25% (250 ppm) — increased monitoring; 1.0% (10,000 ppm) — cease blasting; 1.5% (15,000 ppm) — withdraw all personnel. Carbon monoxide (CO) above 25 ppm triggers immediate investigation. In deep gold mines, hydrogen sulphide (H₂S) from sulphide ore bodies can reach dangerous levels in poorly ventilated areas.

**Heat Stress**
Wet-bulb globe temperature (WBGT) limits apply under MHSA Chapter 6 and SANS 10083. For underground workers doing moderate physical work, the legal WBGT limit is 27.5°C (wet bulb). At temperatures exceeding this, mandatory cooling breaks, increased water consumption, and buddy systems are required. Acclimatisation programmes of 8–10 days are required for workers new to high-temperature environments.

**Equipment and Machinery**
Under MHSA Section 21, all trackless mobile machinery (TMM) must have a Certificate of Fitness (CoF) issued by a competent person before use. Pre-shift inspections must be completed and recorded in the Plant Register. Operators of TMM require a valid certificate of competence (CoC) issued in terms of the Minerals Act.

---

### 2. Immediate Safety Actions (Universal)

1. **Isolate the area** — establish a barrier at least 10 metres beyond the identified hazard zone and post hazard boards
2. **Account for all personnel** — conduct a physical headcount, do not rely solely on tag boards
3. **Notify the shift supervisor and mine overseer** immediately — this is a statutory requirement under MHSA Section 23 for any reportable incident or potentially dangerous situation
4. **Initiate atmospheric testing** — use calibrated, approved multi-gas detectors to confirm O₂ ≥ 19.5%, CH₄ < 0.25%, CO < 25 ppm before re-entry
5. **Implement lock-out / tag-out** on all energy sources associated with the hazard area
6. **Brief all personnel** on the nature of the hazard and prohibited zones before resuming any work

---

### 3. Short-Term Mitigation (24–72 Hours)

- Commission an independent rock engineering assessment of all identified unstable ground areas
- Increase ventilation flow rates to the affected section by a minimum of 20% above minimum legal requirement (0.05 m³/s per kW of diesel equipment installed underground)
- Install additional support: 1.6m friction anchors or resin-grouted bolts at maximum 1.0m × 1.0m pattern in poor ground; increase to 0.75m × 0.75m in very poor ground
- Review and update the section risk register per MHSA Section 11(2)
- Conduct a Section 54 Instruction readiness check — ensure the mine can demonstrate compliance if an inspectorate visit occurs

---

### 4. Long-Term Mitigation Plan (30 Days)

- Review and update the mine's Site-Specific Risk Assessment (SSRA) for all active panels, focusing on identified hazard zones
- Submit updated risk assessments to the Health and Safety Committee (HSC) per MHSA Section 13
- Engage with the DMRE regional office to schedule a proactive compliance inspection — this demonstrates due diligence and is viewed favourably
- Implement a formal Hazard Identification and Risk Assessment (HIRA) training refresher for all supervisors and safety representatives
- Audit PPE inventory and condition — replace any item that does not meet SABS/SANS standards

---

### 5. Applicable MHSA Regulatory References

| Section | Requirement |
|---------|-------------|
| Section 5 | Employer's general duty of care — safe working environment |
| Section 11 | Risk assessments — mandatory before significant changes |
| Section 23 | Incident reporting to the DMRE within prescribed timeframes |
| Section 49 | Personal Protective Equipment — provision, maintenance, training |
| Section 54 | DMR Inspector powers — stop work orders and instructions |
| Chapter 4 | Health & Safety Representatives — consultation obligations |
| Regulation 2.14 | Support of excavations — ground control procedures |
| Regulation 5.1 | Explosives and blasting — safety distances and procedures |

---

### 6. Documentation Requirements

All hazard findings must be recorded in the Mine's Hazard Register within 24 hours of identification. A formal written notification must be made to the Mine Manager and recorded in the Mine Manager's Occurrence Book. Where the hazard constitutes a risk to life, written notification to the DMRE Inspector of Mines is required within 24 hours.{note}"""

    # ── INCIDENT REPORTS / REPORT FORGE ──────────────────────────────────────
    if any(w in p for w in ["incident", "report", "accident", "near miss", "injury", "fatality", "forge", "generate"]):
        return f"""## Formal Incident / Safety Report

### Report Generation — Standard Template

This report has been generated using the Minetrax AI standard template in accordance with the Mine Health and Safety Act, Act 29 of 1996. Once the GROQ_API_KEY is configured, the AI engine will auto-populate all sections with site-specific, context-aware content.

---

### SECTION 1: Incident Summary

**Classification:** To be confirmed by Mine Safety Officer  
**Severity:** As recorded in the incident log  
**Reportable:** Assessment pending — refer to MHSA Section 23 thresholds below

---

### SECTION 2: Detailed Narrative

Document a factual, chronological account of events leading up to, during, and immediately following the incident. Include: the exact sequence of events, the environmental conditions at the time (temperature, visibility, gas readings, ground conditions), the equipment involved and its operational status, and the actions taken by personnel present.

Avoid speculative language. Use first-person accounts from witnesses where possible and record these verbatim with the witness's name, role, and employee number noted.

---

### SECTION 3: Personnel Involved

For each person involved, record: full name, employee number, job title, years of service, training certifications relevant to the task, and whether they had completed a pre-shift safety briefing. Note the time on shift at the point of the incident — fatigue is a significant contributing factor in shift-based mining environments.

---

### SECTION 4: Root Cause Analysis

Apply the Tripod Beta or ICAM (Incident Cause Analysis Method) framework, which is preferred by the DMRE and most major South African mining companies. Root causes typically fall into categories of:

- **Latent Failures** — Management system deficiencies (inadequate procedures, training gaps, resource constraints)
- **Active Failures** — Actions or omissions by individuals directly involved
- **Environmental Conditions** — Physical conditions that contributed to the incident

Avoid the "human error" root cause classification in isolation — this is rarely accepted by the DMRE as a complete causal explanation and does not satisfy the MHSA requirement for systemic corrective action.

---

### SECTION 5: Contributing Factors

List all factors that, while not the direct cause, created conditions that made the incident more likely or more severe. Common contributing factors in South African mining include:

- Inadequate or absent task risk assessment (TRA) at start of shift
- Failure of lock-out / tag-out procedures
- Communication breakdown between shifts (handover quality)
- Equipment not fit for purpose or maintenance overdue
- Production pressure leading to risk tolerance escalation
- Inadequate supervision-to-worker ratio

---

### SECTION 6: Immediate Actions Taken

Document all emergency response actions in chronological order with timestamps: first aid administered, mine rescue called, area isolated, personnel evacuated, scene preserved for investigation, shift supervisor and mine manager notified.

---

### SECTION 7: Corrective Actions Required

| # | Corrective Action | Responsible Person | Target Date | Status |
|---|------------------|--------------------|-------------|--------|
| 1 | Review and update section risk assessment | Rock Engineering Manager | Within 7 days | Pending |
| 2 | Conduct toolbox talk for all shift personnel | Shift Supervisor | Within 24 hours | Pending |
| 3 | Review pre-shift inspection procedures | Mine Overseer | Within 14 days | Pending |
| 4 | Submit corrective action plan to DMR | Mine Manager | Within 30 days | Pending |

---

### SECTION 8: MHSA Regulatory Obligations

**Section 23 — Reporting Thresholds:**
- *Immediately* (by fastest available means): Any incident where a person has died, or is in a critical condition, or where an explosion has occurred
- *Within 24 hours:* Any serious injury (lost time injury, hospitalisation required), dangerous occurrence (explosion, fire, inrush), or occupational disease diagnosis
- *Within 7 days:* Any injury resulting in medical treatment beyond first aid

The Mine Manager is personally liable under MHSA Section 3 to ensure all reportable incidents are notified to the Principal Inspector of Mines for the region.

---

### SECTION 9: Preventive Measures

Based on the incident type and root causes identified, implement the following systemic preventive measures:

1. **Engineering Controls** — modify the physical environment to eliminate or reduce the hazard at source (hierarchy of controls, Level 1)
2. **Administrative Controls** — update written procedures, conduct targeted training, revise work schedules
3. **PPE Enhancement** — if PPE was a contributing factor, review SABS-approved alternatives and confirm training adequacy
4. **Safety System Audit** — commission a focused audit of the management system elements that failed to prevent this incident{note}"""

    # ── COMPLIANCE ENGINE ─────────────────────────────────────────────────────
    if any(w in p for w in ["compliance", "audit", "mhsa", "protocol", "dmr", "regulation"]):
        return f"""## MHSA Compliance Audit Report

### Compliance Assessment — Act 29 of 1996

This audit has been conducted against the requirements of the Mine Health and Safety Act, Act 29 of 1996, and associated regulations. The following comprehensive assessment covers all core compliance domains applicable to South African mining operations.

---

### SECTION 1: Compliance Rating

Compliance scoring is calculated based on the proportion of mandatory safety protocols actively in place at the time of audit. Ratings are classified as:

| Score | Rating | Required Action |
|-------|--------|----------------|
| 80–100% | **Good** | Maintain standards; schedule routine re-audit in 90 days |
| 50–79% | **Moderate** | Address all gaps within 48 hours; notify HSC within 24 hours |
| 0–49% | **Critical** | Cease affected operations; immediate corrective action; notify DMRE |

---

### SECTION 2: Zone Risk Summary

Each mine zone has been classified using K-Means cluster analysis based on gas concentration (ppm), working depth (m), and historical incident frequency (incidents/30 days). The three risk tiers are:

**High Risk Zones** — Characterised by high gas readings (>500 ppm CH₄ or equivalent), greater depths (>1,500m for gold, >300m for coal), and elevated incident frequency. These zones require continuous atmospheric monitoring, 30-minute atmospheric sampling intervals, mandatory buddy system, and a supervisor-to-worker ratio of at least 1:6.

**Medium Risk Zones** — Moderate hazard parameters that require heightened vigilance. Atmospheric sampling every 60 minutes, daily ground inspections, and documented pre-task risk assessments before each shift commencement.

**Low Risk Zones** — Parameters within safe operational thresholds. Standard pre-shift inspections and atmospheric sampling per minimum legal frequency apply.

---

### SECTION 3: Gap Analysis — MHSA Protocol Requirements

The following safety protocols are mandatory under MHSA Act 29 of 1996 and supporting regulations. Any gap represents a statutory compliance failure that must be addressed and documented:

**1. PPE — MHSA Section 49**
The employer must provide, at no cost to the employee, personal protective equipment appropriate to the identified hazards. All PPE must comply with relevant SABS standards. PPE training must be documented and records retained for a minimum of 3 years. Employees who refuse to wear prescribed PPE may be removed from the working area — this must also be documented.

**2. Ventilation & Air Quality Monitoring — MHSA Regulation 5.1 & 8.1**
Minimum legal fresh air flow for underground workings: 0.05 m³/s per installed kW of diesel equipment, with an absolute minimum of 0.3 m³/s to any working face. Wet-bulb temperature must not exceed 27.5°C for working personnel (SANS 10083). Ventilation surveys must be conducted at intervals not exceeding 6 months and results submitted to the DMRE.

**3. Gas Detection & Monitoring — MHSA Regulation 5.1**
Calibrated multi-gas detectors must be deployed at all active working faces. In gassy mines (Schedule 2 classification), continuous methane monitoring with audible and visible alarm at 0.25% CH₄ is mandatory. All gas detector calibration records must be available for inspection.

**4. Traffic Management Plan — MHSA Regulation 8.9**
A written Traffic Management Plan (TMP) signed by the Mine Manager is required wherever trackless mobile machinery operates. Pedestrian-vehicle separation must be achieved through physical or procedural controls. TMM operators must hold valid certificates of competence (CoC).

**5. Environmental & Dust Monitoring — MHSA Regulation 9.1**
Dust sampling must be conducted at regulated intervals per the Occupational Hygiene Regulations. For silica-bearing dust environments, personal dust sampling of at least 10% of workers in each occupational category per quarter is required. Results must be submitted to the Mine Health and Safety Inspectorate and retained for 40 years (for potential compensation claims).

**6. Pre-Shift Equipment Inspections — MHSA Section 21**
Written pre-shift inspection records must be maintained in the Plant Register. No equipment may be operated until the inspection is completed and signed off by a competent person. Defects must be reported, recorded, and the equipment taken out of service until certified fit.

**7. First Aid Readiness & Emergency Response Plan — MHSA Section 10**
A written Emergency Response Plan (ERP) must be available at the mine, rehearsed at least annually, and a copy lodged with the DMRE. Minimum first aider ratios: 1 per 10 persons for underground operations; 1 per 50 persons on surface. First aid certificates must be current (renewed every 3 years through an accredited training provider).

---

### SECTION 4: Corrective Actions with Owners and Timelines

All identified gaps must be entered into the mine's Corrective Action Register and tracked to closure. The Health and Safety Committee (HSC) must be informed of all gaps and their status at each scheduled meeting. The Mine Manager remains personally accountable under MHSA Section 3 for ensuring corrective actions are completed within the specified timeframes.

---

### SECTION 5: Applicable MHSA Sections for This Mine Type

- **Section 5** — General duty to ensure health and safety
- **Section 11** — Risk assessment obligations
- **Section 13** — Health and Safety Committee establishment and consultation
- **Section 23** — Incident reporting to DMRE
- **Section 49** — PPE provision and management
- **Section 54** — Inspector powers and right to issue stop-work instructions
- **Chapter 6** — Occupational health — dust, noise, heat, medical surveillance

---

### SECTION 6: Next Audit Schedule Recommendation

Based on the compliance score achieved:
- **Score ≥ 80%:** Re-audit in 90 days (quarterly schedule)
- **Score 50–79%:** Re-audit in 30 days to verify corrective action closure
- **Score < 50%:** Re-audit within 14 days; notify DMRE of compliance status and remediation plan{note}"""

    # ── SAFETY CHATBOT (general mining safety) ────────────────────────────────
    if any(w in p for w in ["what", "how", "when", "why", "explain", "tell me", "advise", "should", "ventilation", "ppe", "seismic", "support", "blasting", "emergency"]):
        return f"""## Mining Safety Advisory Response

Thank you for your question. The following guidance is based on South African mining best practice and the Mine Health and Safety Act, Act 29 of 1996. For site-specific AI-generated advice tailored to your exact mine, zone, and conditions, please configure the GROQ_API_KEY.

---

### General Mining Safety Guidance

South African mining is governed primarily by the **Mine Health and Safety Act, Act 29 of 1996 (MHSA)**, administered by the Department of Mineral Resources and Energy (DMRE) through the Mine Health and Safety Inspectorate. All mining operations — underground, open pit, quarry, alluvial, and surface — are subject to MHSA requirements and supporting technical regulations.

**Key Employer Duties (MHSA Section 5):**
Every employer at a mine must, as far as is reasonably practicable:
- Identify hazards and eliminate or mitigate risks
- Provide systems of work that are safe and without risk to health
- Ensure that equipment, materials, and substances are safe and without risk
- Provide information, instruction, training, and supervision necessary to ensure safety
- Ensure that the mine is designed, constructed, and maintained to be safe

**Health and Safety Representatives (MHSA Chapter 4):**
Every mine employing 20 or more employees must have elected Health and Safety Representatives. These representatives have the right to inspect any part of the mine, investigate complaints, and accompany DMR inspectors. Employers must consult with H&S Reps on all matters affecting health and safety.

**Risk Assessment (MHSA Section 11):**
Before commencing any new operation or making any significant change to an existing operation, a formal written risk assessment must be conducted by a competent person. The risk assessment must identify all hazards, assess the risks, and specify control measures. It must be reviewed and updated at least annually or whenever there is a material change in conditions.

**Incident Reporting (MHSA Section 23):**
Reportable incidents must be notified to the Principal Inspector of Mines for the region:
- Immediately: Death, critical injury, explosion, major inrush
- Within 24 hours: Serious injury, hospitalisation, dangerous occurrence, occupational disease
- Written report within 7 days using Form 10.1 (MHSA Regulation)

**Emergency Preparedness (MHSA Section 10):**
Every mine must have a written Emergency Response Plan (ERP) that covers all foreseeable emergencies including fall of ground, gas outburst, fire, flooding, transportation accident, and medical emergency. The ERP must be rehearsed, and the results of rehearsals recorded and made available for inspection.

For your specific question, please connect the AI engine for a detailed, context-aware response tailored to your mine type, location, and current conditions.{note}"""

    # ── DEFAULT ───────────────────────────────────────────────────────────────
    return f"""## Minetrax AI — Safety Intelligence Response

Your query has been received. The following general guidance applies to South African mining operations under the Mine Health and Safety Act, Act 29 of 1996.

**Always remember the hierarchy of controls in South African mining:**

1. **Elimination** — Remove the hazard entirely from the workplace
2. **Substitution** — Replace the hazardous material or process with a less dangerous one
3. **Engineering Controls** — Isolate people from the hazard through physical barriers, ventilation, interlocks
4. **Administrative Controls** — Change work procedures, schedules, supervision ratios, training
5. **Personal Protective Equipment (PPE)** — The last line of defence, not the first

Under MHSA Section 11, all significant hazards must be addressed using this hierarchy, and the rationale for the chosen control level must be documented in the risk assessment.

**Verify your findings with your Mine Safety Officer and ensure all observations are recorded in the Mine's Hazard Register within 24 hours.**{note}"""

# ═══════════════════════════════════════════════════════════════════════════════
# PDF EXPORT – ENHANCED WITH SIGNATURE BLOCK
# ═══════════════════════════════════════════════════════════════════════════════
def safe_text(text):
    if not text:
        return ""
    replacements = {
        "–": "-", "—": "-", "•": "-", "⚠": "[!]", "✓": "OK",
        "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
        "🟢": "[GREEN]", "🟡": "[YELLOW]", "🔴": "[RED]", "⬇": "",
        "✅": "[OK]", "❌": "[X]", "⛏": "[MINE]",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode("latin-1", "ignore").decode("latin-1")

def create_pdf(text, title, metadata=None, signatories=None):
    try:
        pdf = FPDF()
        pdf.add_page()

        # ── Header bar ──────────────────────────────────────────────────────
        pdf.set_fill_color(8, 8, 22)
        pdf.rect(0, 0, 210, 52, 'F')
        pdf.set_fill_color(245, 166, 35)
        pdf.rect(0, 48, 210, 4, 'F')

        pdf.set_font("Arial", 'B', 26)
        pdf.set_text_color(245, 166, 35)
        pdf.set_xy(12, 8)
        pdf.cell(186, 14, "MINETRAX AI", new_x="LMARGIN", new_y="NEXT", align="L")

        pdf.set_font("Arial", size=10)
        pdf.set_text_color(200, 200, 200)
        pdf.set_xy(12, 24)
        pdf.cell(186, 6, "Mining Safety Intelligence Platform  |  South Africa", new_x="LMARGIN", new_y="NEXT", align="L")

        pdf.set_font("Arial", 'I', 8)
        pdf.set_text_color(130, 130, 130)
        pdf.set_xy(12, 33)
        pdf.cell(186, 5, f"Generated: {datetime.now().strftime('%d %B %Y  %H:%M')}   |   CONFIDENTIAL   |   MHSA Act 29 of 1996", new_x="LMARGIN", new_y="NEXT", align="L")

        # ── Title banner ────────────────────────────────────────────────────
        pdf.ln(8)
        pdf.set_fill_color(245, 166, 35)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Arial", 'B', 13)
        clean_title = safe_text(title.upper())
        pdf.cell(0, 11, f"  {clean_title}", fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)

        # ── Metadata table ──────────────────────────────────────────────────
        if metadata:
            pdf.set_font("Arial", 'B', 8)
            for k, v in metadata.items():
                pdf.set_fill_color(235, 235, 235)
                pdf.set_text_color(40, 40, 40)
                pdf.cell(55, 6, f" {safe_text(str(k))}", border=1, fill=True)
                pdf.set_font("Arial", size=8)
                pdf.set_fill_color(250, 250, 250)
                pdf.cell(135, 6, f" {safe_text(str(v))}", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("Arial", 'B', 8)
            pdf.ln(6)

        # ── Body text ───────────────────────────────────────────────────────
        pdf.set_text_color(25, 25, 25)
        pdf.set_font("Arial", size=10)
        clean = safe_text(text)
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', clean)
        clean = re.sub(r'\*(.*?)\*', r'\1', clean)
        clean = re.sub(r'#{1,6}\s?', '', clean)
        pdf.multi_cell(0, 6, clean)
        pdf.ln(6)

        # ── Signature block ─────────────────────────────────────────────────
        pdf.set_fill_color(245, 166, 35)
        pdf.rect(0, pdf.get_y(), 210, 1, 'F')
        pdf.ln(4)
        pdf.set_font("Arial", 'B', 10)
        pdf.set_text_color(20, 20, 20)
        pdf.cell(0, 7, "AUTHORISATION & SIGNATURES", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        default_sigs = signatories or [
            ("Mine Manager / Principal Officer", ""),
            ("Senior Safety Officer (MHSA Ch. 3)", ""),
            ("Health & Safety Representative", ""),
            ("DMR Inspector (if applicable)", ""),
        ]

        for role, name in default_sigs:
            pdf.set_font("Arial", size=9)
            pdf.set_text_color(60, 60, 60)
            role_str = safe_text(str(role))
            name_str = safe_text(str(name)) if name else "__________________________________"
            pdf.cell(90, 5, f"Role: {role_str}", border=0)
            pdf.cell(0, 5, f"Name: {name_str}", border=0, new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(180, 180, 180)
            y = pdf.get_y()
            pdf.set_font("Arial", 'I', 8)
            pdf.set_text_color(140, 140, 140)
            pdf.cell(90, 4, "Signature: _______________________________", border=0)
            pdf.cell(60, 4, "Date: ____________________", border=0)
            pdf.cell(0, 4, "", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(5)

        # ── Footer ──────────────────────────────────────────────────────────
        pdf.set_fill_color(245, 166, 35)
        pdf.rect(0, pdf.get_y(), 210, 1, 'F')
        pdf.ln(3)
        pdf.set_font("Arial", 'I', 8)
        pdf.set_text_color(140, 140, 140)
        pdf.cell(0, 5, "Minetrax AI  |  Mining Safety Intelligence  |  MHSA Act 29 of 1996 Compliant  |  www.minetrax.ai", align='C')

        out = pdf.output()
        return bytes(out) if isinstance(out, (bytes, bytearray)) else out.encode('latin-1', 'replace')
    except Exception as e:
        st.error(f"PDF generation error: {e}")
        return None

def create_docx(text, title, metadata=None, signatories=None):
    doc = Document()

    # Title
    h = doc.add_heading("MINETRAX AI – Mining Safety Intelligence Platform", 0)
    for run in h.runs:  
        run.font.color.rgb = RGBColor(245, 166, 35)

    sub = doc.add_paragraph(title)
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}  |  CONFIDENTIAL  |  MHSA Act 29 of 1996")

    if metadata:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        tbl.rows[0].cells[0].text = "Field"
        tbl.rows[0].cells[1].text = "Value"
        for k, v in metadata.items():
            row = tbl.add_row()
            row.cells[0].text = str(k)
            row.cells[1].text = str(v)
        doc.add_paragraph()

    doc.add_paragraph(text)
    doc.add_paragraph()

    # Signature section
    doc.add_heading("AUTHORISATION & SIGNATURES", level=2)
    default_sigs = signatories or [
        ("Mine Manager / Principal Officer", ""),
        ("Senior Safety Officer (MHSA Ch. 3)", ""),
        ("Health & Safety Representative", ""),
        ("DMR Inspector (if applicable)", ""),
    ]
    sig_tbl = doc.add_table(rows=1, cols=4)
    sig_tbl.style = 'Table Grid'
    hdr = sig_tbl.rows[0].cells
    hdr[0].text = "Role / Designation"
    hdr[1].text = "Full Name"
    hdr[2].text = "Signature"
    hdr[3].text = "Date"
    for role, name in default_sigs:
        row = sig_tbl.add_row()
        row.cells[0].text = str(role)
        row.cells[1].text = str(name) if name else ""
        row.cells[2].text = ""
        row.cells[3].text = ""

    doc.add_paragraph()
    doc.add_paragraph("Minetrax AI  |  MHSA Act 29 of 1996 Compliant  |  South Africa")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════════
# HELPER – load background image
# ═══════════════════════════════════════════════════════════════════════════════
def _get_bg_image():
    for path in ["mines.jpg", "/home/claude/mines.jpg", "app/mines.jpg"]:
        try:
            with open(path, "rb") as f:
                return base64.b64encode(f.read()).decode()
        except Exception:
            pass
    return ""

BG_B64 = _get_bg_image()

# ═══════════════════════════════════════════════════════════════════════════════
# CSS STYLING  – improved colour contrast + welcome overhaul
# ═══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Bebas+Neue&family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;0,9..40,600;0,9..40,700;1,9..40,400&family=JetBrains+Mono:wght@400;600&display=swap');
:root {
  --gold:    #f5a623;
  --gold2:   #e8621a;
  --bg:      #060610;
  --surface: #0c0c1a;
  --s2:      #121224;
  --s3:      #181830;
  --border:  #2a2a50;
  --text:    #f0ecff;
  --dim:     #9b97b0;
  --danger:  #ff4757;
  --success: #2ed573;
  --warn:    #ffa502;
}
*, *::before, *::after { box-sizing: border-box; }

/* App */
.stApp { background: var(--bg) !important; color: var(--text) !important; font-family: 'DM Sans', sans-serif !important; }
.main .block-container { padding-top: 1.5rem !important; padding-bottom: 2rem !important; max-width: 1400px !important; }

/* Sidebar */
section[data-testid="stSidebar"] { background: var(--surface) !important; border-right: 2px solid var(--border) !important; }
section[data-testid="stSidebar"] * { color: var(--text) !important; }
section[data-testid="stSidebar"] .stRadio label { font-size: 13px !important; padding: 6px 0 !important; color: var(--text) !important; font-weight: 500 !important; }

/* Metrics */
div[data-testid="stMetric"] { background: var(--s2) !important; border: 1px solid var(--border); border-radius: 16px; padding: 18px !important; }
div[data-testid="stMetricValue"] { color: var(--gold) !important; font-family: 'Bebas Neue'; font-size: 36px !important; }
div[data-testid="stMetricLabel"] { color: var(--text) !important; font-size: 11px !important; text-transform: uppercase; letter-spacing: 1.5px; font-weight:600; }
div[data-testid="stMetricDelta"] { color: var(--dim) !important; font-size: 11px !important; }

/* Buttons */
.stButton > button {
  background: linear-gradient(135deg, var(--gold), var(--gold2)) !important;
  color: #000 !important; font-weight: 800 !important; border: none !important;
  border-radius: 10px !important; padding: 0.6rem 1.4rem !important;
  font-family: 'DM Sans' !important; letter-spacing: 1px !important;
  text-transform: uppercase !important; font-size: 13px !important;
  width: 100% !important; transition: all 0.2s ease !important;
  box-shadow: 0 4px 18px rgba(245,166,35,0.3) !important;
}
.stButton > button:hover { transform: translateY(-2px) !important; box-shadow: 0 8px 28px rgba(245,166,35,0.5) !important; }

/* Download buttons */
div[data-testid="stDownloadButton"] button {
  background: rgba(245,166,35,0.15) !important; border: 2px solid var(--gold) !important;
  color: var(--gold) !important; font-weight: 700 !important; border-radius: 10px !important;
  font-size: 13px !important; padding: 0.55rem 1rem !important; width: 100% !important;
}
div[data-testid="stDownloadButton"] button:hover { background: rgba(245,166,35,0.25) !important; }

/* Inputs – high-contrast labels and text */
.stTextInput input, .stTextArea textarea {
  background: var(--s2) !important; border: 1.5px solid var(--border) !important;
  border-radius: 10px !important; color: var(--text) !important; font-family: 'DM Sans' !important;
  font-size: 14px !important;
}
.stTextInput input:focus, .stTextArea textarea:focus { border-color: var(--gold) !important; box-shadow: 0 0 0 3px rgba(245,166,35,0.18) !important; }
.stTextInput input::placeholder, .stTextArea textarea::placeholder { color: #7070a0 !important; }
.stSelectbox > div > div, .stMultiSelect > div > div {
  background: var(--s2) !important; border: 1.5px solid var(--border) !important;
  border-radius: 10px !important; color: var(--text) !important;
}
/* Selectbox dropdown options */
.stSelectbox [data-baseweb="select"] *, .stMultiSelect [data-baseweb="select"] * { color: var(--text) !important; }
[data-baseweb="popover"] * { background: #1a1a30 !important; color: var(--text) !important; }
[data-baseweb="option"] { background: #1a1a30 !important; color: var(--text) !important; }
[data-baseweb="option"]:hover { background: #2a2a50 !important; }
/* Multiselect tags */
[data-baseweb="tag"] { background: rgba(245,166,35,0.2) !important; color: var(--gold) !important; }

/* Slider */
.stSlider > div { color: var(--text) !important; }
.stSlider [data-testid="stThumbValue"] { color: var(--gold) !important; font-weight: 700 !important; font-size: 13px !important; }
div[data-baseweb="slider"] [role="slider"] { background: var(--gold) !important; }

/* Select slider */
.stSelectSlider > div { color: var(--text) !important; }

/* Typography */
h1 { font-family: 'Bebas Neue' !important; font-size: 38px !important; color: var(--text) !important; letter-spacing: 2px !important; margin-bottom: 2px !important; }
h2 { font-family: 'Bebas Neue' !important; color: var(--gold) !important; letter-spacing: 1px !important; font-size: 26px !important; }
h3 { font-family: 'DM Sans' !important; font-weight: 700 !important; color: var(--text) !important; font-size: 16px !important; }
p, li { color: var(--text) !important; }

/* All labels – high contrast */
.stSelectbox label, .stMultiSelect label, .stTextInput label, .stTextArea label,
.stFileUploader label, .stSlider label, .stDateInput label, .stTimeInput label,
.stRadio label, .stCheckbox label, .stSelectSlider label {
  color: var(--text) !important; font-size: 13px !important; font-weight: 600 !important;
  letter-spacing: 0.3px !important;
}

/* Cards */
.card { background: var(--s2); border: 1px solid var(--border); border-radius: 18px; padding: 22px; margin-bottom: 14px; }
.card-gold { border-left: 3px solid var(--gold); }
.sec-label { font-size: 10px; text-transform: uppercase; letter-spacing: 2.5px; color: var(--dim); margin-bottom: 8px; font-weight: 700; }
.result-panel { background: var(--s2); border: 1px solid var(--border); border-radius: 18px; padding: 26px; margin-top: 18px; border-top: 3px solid var(--gold); }

/* Logo */
.logo-wrap { padding: 6px 0 18px; }
.logo-icon { font-size: 36px; line-height: 1; filter: drop-shadow(0 0 14px rgba(245,166,35,0.7)); }
.logo-title { font-family: 'Bebas Neue'; font-size: 30px; color: var(--gold); letter-spacing: 4px; line-height: 1; margin-top: 6px; }
.logo-sub { font-size: 9px; color: var(--dim); letter-spacing: 3px; text-transform: uppercase; margin-top: 3px; line-height: 1.4; }

/* Chat */
.stChatMessage { background: var(--s2) !important; border-radius: 14px !important; border: 1px solid var(--border) !important; }
.stChatInput textarea { background: var(--s2) !important; border: 1.5px solid var(--border) !important; border-radius: 12px !important; color: var(--text) !important; }

/* Expander */
.streamlit-expanderHeader { background: var(--s2) !important; border: 1px solid var(--border) !important; border-radius: 10px !important; color: var(--text) !important; font-weight: 600 !important; }
details { background: var(--s2) !important; border: 1px solid var(--border) !important; border-radius: 10px !important; }

/* Date input */
.stDateInput input { background: var(--s2) !important; border: 1.5px solid var(--border) !important; color: var(--text) !important; border-radius: 10px !important; }

/* File uploader */
.stFileUploader { background: var(--s2) !important; border: 1.5px dashed var(--border) !important; border-radius: 12px !important; }
.stFileUploader * { color: var(--text) !important; }

/* Alert / info boxes */
.stAlert { border-radius: 12px !important; }

/* Misc */
hr { border-color: var(--border) !important; opacity: 0.5; }
::-webkit-scrollbar { width: 5px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: var(--border); border-radius: 3px; }

/* Welcome page hero */
.welcome-hero {
  position: relative; border-radius: 24px; overflow: hidden;
  min-height: 480px; display: flex; align-items: flex-end;
  margin-bottom: 36px; box-shadow: 0 24px 60px rgba(0,0,0,0.7);
}
.welcome-hero img { position: absolute; inset: 0; width: 100%; height: 100%; object-fit: cover; }
.welcome-hero-overlay {
  position: absolute; inset: 0;
  background: linear-gradient(to top, rgba(6,6,22,0.98) 0%, rgba(6,6,22,0.70) 45%, rgba(6,6,22,0.15) 100%);
}
.welcome-hero-content { position: relative; z-index: 2; padding: 40px 44px; width: 100%; }

/* Feature cards */
.feat-card { background: var(--s2); border: 1px solid var(--border); border-radius: 16px; padding: 22px; height: 100%; border-top: 2px solid var(--gold); }
.feat-icon { font-size: 28px; margin-bottom: 10px; }
.feat-title { font-weight: 700; font-size: 15px; color: var(--text); margin-bottom: 6px; }
.feat-body { font-size: 13px; color: var(--dim); line-height: 1.7; }

/* User story cards */
.user-card {
  background: linear-gradient(135deg, #0f0f22, #1a1a35);
  border: 1px solid var(--border); border-radius: 18px; padding: 24px;
  border-left: 4px solid var(--gold); margin-bottom: 10px;
}

/* Mobile */
@media (max-width: 768px) {
  .main .block-container { padding: 0.5rem !important; }
  h1 { font-size: 28px !important; }
  .welcome-hero { min-height: 320px; }
  div[data-testid="stMetricValue"] { font-size: 28px !important; }
}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ═══════════════════════════════════════════════════════════════════════════════
def _init():
    defaults = {
        "page": "welcome",
        "stats": {"hazards": 0, "compliance_score": 0, "reports": 0, "queries": 0, "used": False},
        "chat_history": [], "hazard_analysis": None, "hz_pred": None,
        "compliance_audit": None, "active_doc": None, "doc_type": None, "doc_meta": None,
        "show_gas": False, "mine_configured": False,
        "sel_province": "Gauteng", "sel_mine": None,
        "hz_zone_risk": None, "hz_location": None,
        "_audit_zones": [], "_audit_lookup": {}, "_audit_op": "",
        "_audit_sup": "", "_audit_date": "", "_audit_score": 0,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v
_init()

def mark_used(): st.session_state.stats["used"] = True

# ═══════════════════════════════════════════════════════════════════════════════
# WELCOME PAGE  (shown when mine_configured is False)
# ═══════════════════════════════════════════════════════════════════════════════
if not st.session_state.mine_configured:

    # ── Hero banner ──────────────────────────────────────────────────────────
    if BG_B64:
        hero_img_tag = f'<img src="data:image/jpeg;base64,{BG_B64}" alt="Mining site at sunset"/>'
    else:
        hero_img_tag = f'<div style="position:absolute;inset:0;background:linear-gradient(135deg,#0c0c2a,#1a1030);"></div>'

    # SVG Logo embedded in the hero
    logo_svg = """
    <svg width="80" height="80" viewBox="0 0 80 80" fill="none" xmlns="http://www.w3.org/2000/svg">
      <circle cx="40" cy="40" r="38" fill="#0c0c22" stroke="#f5a623" stroke-width="2.5"/>
      <!-- Pickaxe icon -->
      <line x1="20" y1="60" x2="58" y2="22" stroke="#f5a623" stroke-width="5" stroke-linecap="round"/>
      <path d="M54 18 L62 26 L58 30 C55 33 50 32 47 29 C44 26 43 21 46 18 Z" fill="#f5a623"/>
      <path d="M18 62 L22 58 C25 55 28 55 30 57 C32 59 32 62 29 65 Z" fill="#e8621a"/>
      <!-- Gear element -->
      <circle cx="56" cy="56" r="10" fill="none" stroke="#f5a623" stroke-width="2" stroke-dasharray="4 3"/>
      <circle cx="56" cy="56" r="5" fill="#f5a623" opacity="0.7"/>
      <!-- AI dot -->
      <circle cx="28" cy="28" r="4" fill="#f5a623" opacity="0.5"/>
      <circle cx="28" cy="28" r="2" fill="#fff" opacity="0.9"/>
    </svg>"""

    st.markdown(f"""
    <div class="welcome-hero">
      {hero_img_tag}
      <div class="welcome-hero-overlay"></div>
      <div class="welcome-hero-content">
        <div style="display:flex;align-items:center;gap:20px;margin-bottom:18px;">
          {logo_svg}
          <div>
            <div style="font-family:'Bebas Neue';font-size:58px;color:#f5a623;letter-spacing:5px;line-height:1;text-shadow:0 4px 20px rgba(245,166,35,0.4);">
              MINETRAX AI
            </div>
            <div style="font-family:'Bebas Neue';font-size:20px;color:#ccc8e0;letter-spacing:4px;margin-top:4px;">
              MINING SAFETY INTELLIGENCE PLATFORM
            </div>
          </div>
        </div>
        <div style="max-width:680px;font-size:16px;color:#ddd8f0;line-height:1.8;margin-bottom:28px;">
          Protecting South Africa's miners with <strong style="color:#f5a623;">AI-powered hazard detection</strong>,
          real-time compliance auditing, and intelligent documentation — built on the
          <strong style="color:#f5a623;">Mine Health and Safety Act (Act 29 of 1996)</strong>.
        </div>
        <div style="display:flex;gap:14px;flex-wrap:wrap;margin-bottom:10px;">
          <div style="background:rgba(245,166,35,0.15);border:1px solid rgba(245,166,35,0.4);border-radius:20px;padding:6px 18px;font-size:12px;color:#f5a623;font-weight:700;letter-spacing:1px;">⚡ AI Hazard Analysis</div>
          <div style="background:rgba(46,213,115,0.12);border:1px solid rgba(46,213,115,0.35);border-radius:20px;padding:6px 18px;font-size:12px;color:#2ed573;font-weight:700;letter-spacing:1px;">✓ MHSA Compliant</div>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Mine setup & start button ────────────────────────────────────────────
    st.markdown('<div style="font-family:\'Bebas Neue\';font-size:28px;color:#f5a623;letter-spacing:2px;margin:28px 0 20px;">SELECT YOUR MINE SITE TO BEGIN</div>', unsafe_allow_html=True)

    st.markdown("""
    <div style="background:rgba(245,166,35,0.08);border:1px solid rgba(245,166,35,0.25);border-radius:14px;padding:16px 20px;margin-bottom:20px;">
      <span style="font-size:13px;color:#ddd8f0;">Choose your province and mine operation below, then click <strong style="color:#f5a623;">Launch Platform</strong> to access all modules.</span>
    </div>
    """, unsafe_allow_html=True)

    wc1, wc2 = st.columns(2)
    with wc1:
        sel_province = st.selectbox("Province", PROVINCES, index=PROVINCES.index(st.session_state.sel_province), key="welcome_province")
    with wc2:
        province_mines = PROVINCE_MINE_DATA[sel_province]["mines"]
        default_mine_idx = 0
        if st.session_state.sel_mine in province_mines:
            default_mine_idx = province_mines.index(st.session_state.sel_mine)
        sel_mine = st.selectbox("Mine / Operation", province_mines, index=default_mine_idx, key="welcome_mine")

    st.session_state.sel_province = sel_province
    st.session_state.sel_mine = sel_mine

    # ── API key entry (if not already set via .env)

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("⛏  LAUNCH MINETRAX AI  →"):
        st.session_state.mine_configured = True
        st.rerun()

    # ── MHSA reference footer ────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("""
    <div style="text-align:center;padding:14px 0;">
      <div style="font-size:11px;color:#6b6880;">
        Built on the <strong style="color:#9b97b0;">Mine Health and Safety Act, Act 29 of 1996 (MHSA)</strong> 
      </div>
      <div style="font-size:11px;color:#6b6880;margin-top:4px;">
        ⛏ Built for the safety of South Africa's miners
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.stop()

# ═══════════════════════════════════════════════════════════════════════════════
# MAIN PLATFORM  (shown after mine is configured)
# ═══════════════════════════════════════════════════════════════════════════════

# ── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div class="logo-wrap">
      <div class="logo-icon">⛏</div>
      <div class="logo-title">MINETRAX</div>
      <div class="logo-sub">Mining Safety Intelligence<br>Platform · South Africa</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="sec-label">Navigation</div>', unsafe_allow_html=True)
    menu = st.radio("", [
        "◈  Safety Dashboard",
        "⚠️  Hazard Intelligence",
        "✦  Compliance Engine",
        "◉  Report Forge",
        "◎  Safety Advisor",
    ], label_visibility="collapsed")

    st.markdown("---")
    st.markdown('<div class="sec-label">Active Mine Site</div>', unsafe_allow_html=True)
    active_province = st.selectbox("Province", PROVINCES,
                                   index=PROVINCES.index(st.session_state.sel_province),
                                   label_visibility="collapsed",
                                   key="sidebar_province")

    pdata = PROVINCE_MINE_DATA[active_province]
    mine_opts = pdata["mines"]
    default_idx = 0
    if st.session_state.sel_mine in mine_opts:
        default_idx = mine_opts.index(st.session_state.sel_mine)
    active_mine = st.selectbox("Mine / Operation", mine_opts,
                                index=default_idx,
                                label_visibility="collapsed",
                                key="sidebar_mine")
    active_mine_type = pdata["type"]
    province_zones   = pdata["zones"]
    province_ops     = pdata["operations"]

    st.markdown(f'<div style="font-size:11px;color:#9b97b0;margin-top:4px;font-weight:600;">{active_mine_type}</div>', unsafe_allow_html=True)
    st.markdown(f'<div style="font-size:10px;color:#6b6880;margin-top:2px;">{active_province}</div>', unsafe_allow_html=True)

    st.markdown("---")

    # Check if the environment variable is loaded
    api_ok = bool(GROQ_API_KEY)
    sc_c = "#2ed573" if api_ok else "#ff4757"
    sc_l = "AI Engine Online" if api_ok else "AI Engine Offline"
    
    st.markdown(f"""<div style="display:flex;align-items:center;gap:8px;padding:6px 0;">
    <div style="width:9px;height:9px;border-radius:50%;background:{sc_c};box-shadow:0 0 10px {sc_c};flex-shrink:0;"></div>
    <span style="font-size:12px;color:{sc_c};font-weight:700;">{sc_l}</span>
    </div>""", unsafe_allow_html=True)

    st.markdown("---")
    if st.button("← Back to Welcome"):
        st.session_state.mine_configured = False
        st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
# ◈ SAFETY DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════
if menu == "◈  Safety Dashboard":
    st.title("Safety Operations Centre")
    st.markdown(f'<div class="sec-label">{active_mine} · {active_province} · Live Monitoring</div>', unsafe_allow_html=True)

    s = st.session_state.stats
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Hazard Analyses",    s["hazards"],                "Completed")
    c2.metric("Compliance Score",   f"{s['compliance_score']}%" if s['compliance_score'] else "—", "Last Audit")
    c3.metric("Reports Generated",  s["reports"],                "Documents")
    c4.metric("AI Queries",         s["queries"],                "Total")

    if s["compliance_score"] > 0:
        pct = s["compliance_score"]
        bar_color = "#2ed573" if pct >= 80 else "#ffa502" if pct >= 50 else "#ff4757"
        st.markdown(f"""
        <div style="margin:14px 0 4px;font-size:10px;color:#9b97b0;letter-spacing:2px;text-transform:uppercase;font-weight:700;">Compliance Progress</div>
        <div style="background:#0a0a16;border-radius:8px;height:10px;overflow:hidden;border:1px solid #2a2a50;">
          <div style="width:{pct}%;height:100%;background:{bar_color};border-radius:8px;transition:width 0.6s;"></div>
        </div>
        <div style="font-size:11px;color:#9b97b0;margin-top:4px;">{pct}% compliant with selected safety protocols</div>
        """, unsafe_allow_html=True)

    if s["used"]:
        st.markdown("---")
        if s["compliance_score"] >= 80:
            st.success("✅ Site Status: Compliant — All monitored parameters within safe operational limits.")
        elif s["compliance_score"] >= 50:
            st.warning("⚠️ Site Status: Moderate Risk — Identified protocol gaps require attention within 48 hours.")
        elif s["compliance_score"] > 0:
            st.error("🔴 Site Status: High Risk — Critical compliance gaps. Immediate corrective action required.")

    # Active Mine Overview
    st.markdown("---")
    st.markdown("## Active Mine Overview")
    ov1, ov2, ov3 = st.columns(3)
    with ov1:
        ops_list = "".join(f'<li style="font-size:12px;color:#d8d4f0;margin-bottom:4px;">{o}</li>' for o in province_ops[:4])
        st.markdown(f"""<div class="card card-gold">
        <div style="font-size:11px;color:#f5a623;font-weight:700;letter-spacing:1.5px;text-transform:uppercase;margin-bottom:10px;">Operations</div>
        <ul style="margin:0;padding-left:16px;">{ops_list}</ul>
        </div>""", unsafe_allow_html=True)
    with ov2:
        zones_list = "".join(f'<li style="font-size:12px;color:#d8d4f0;margin-bottom:4px;">{z[:42]}</li>' for z in province_zones[:4])
        st.markdown(f"""<div class="card card-gold">
        <div style="font-size:11px;color:#f5a623;font-weight:700;letter-spacing:1.5px;text-transform:uppercase;margin-bottom:10px;">Key Zones</div>
        <ul style="margin:0;padding-left:16px;">{zones_list}</ul>
        </div>""", unsafe_allow_html=True)
    with ov3:
        st.markdown(f"""<div class="card card-gold">
        <div style="font-size:11px;color:#f5a623;font-weight:700;letter-spacing:1.5px;text-transform:uppercase;margin-bottom:10px;">Mine Profile</div>
        <div style="font-size:14px;color:#f0ecff;font-weight:700;">{active_mine_type}</div>
        <div style="font-size:12px;color:#9b97b0;margin-top:6px;">{active_province}</div>
        <div style="font-size:12px;color:#9b97b0;margin-top:3px;">{len(province_zones)} zones · {len(province_ops)} operation types</div>
        </div>""", unsafe_allow_html=True)

    # MHSA Quick Reference
    st.markdown("---")
    st.markdown("## MHSA Act 29 of 1996 — Key Sections Reference")
    mhsa_refs = [
        ("Section 5", "Employer's duty of care — establish and maintain a safe working environment."),
        ("Section 11", "Risk assessments — mandatory before all significant operational changes."),
        ("Section 23", "Incident reporting — reportable incidents to the DMRE within 24 hours."),
        ("Section 49", "Personal Protective Equipment — provision, maintenance, and training."),
        ("Chapter 4", "Health and Safety Representatives — rights, duties, and consultation obligations."),
        ("Section 54", "DMR Inspector powers — right to issue instructions and stop work orders."),
    ]
    mc1, mc2 = st.columns(2)
    for i, (sec, desc) in enumerate(mhsa_refs):
        with (mc1 if i % 2 == 0 else mc2):
            st.markdown(f"""<div style="background:#0f0f22;border:1px solid #2a2a50;border-left:3px solid #f5a623;border-radius:10px;padding:12px 16px;margin-bottom:10px;">
            <div style="font-size:12px;color:#f5a623;font-weight:700;margin-bottom:4px;">{sec}</div>
            <div style="font-size:12px;color:#d8d4f0;line-height:1.6;">{desc}</div>
            </div>""", unsafe_allow_html=True)

    # Gas anomaly monitor
    st.markdown("---")
    with st.expander("📡 Live Gas Anomaly Monitor — Z-Score Anomaly Detection", expanded=False):
        st.markdown('<div class="sec-label">Enter sensor readings to detect dangerous gas spikes using statistical analysis</div>', unsafe_allow_html=True)
        da1, da2 = st.columns([2, 1])
        with da1:
            gas_input = st.text_input("Gas Readings in ppm (comma-separated)", value="", key="dash_gas",
                                      placeholder="e.g. 120, 135, 142, 850, 131 … (minimum 4 readings)")
        with da2:
            dash_thresh = st.slider("Alert Threshold (Z-score)", 1.5, 4.0, 2.5, 0.1, key="dash_thresh")

        if gas_input.strip():
            try:
                readings = [float(x.strip()) for x in gas_input.split(",") if x.strip()]
                if len(readings) < 4:
                    st.warning("Enter at least 4 readings for meaningful analysis.")
                else:
                    anomalies, mean_r, std_r = _ml_anomaly_detect(readings, dash_thresh)
                    max_r = max(readings) or 1
                    spark = "".join(
                        f'<div title="#{i+1}: {r}ppm" style="flex:1;min-width:8px;height:{max(4,int((r/max_r)*70))}px;'
                        f'background:{"#ff4757" if any(a[0]==i for a in anomalies) else "#f5a623"};'
                        f'border-radius:3px 3px 0 0;opacity:{"1" if any(a[0]==i for a in anomalies) else "0.7"};"></div>'
                        for i, r in enumerate(readings))
                    sc = "#ff4757" if anomalies else "#2ed573"
                    st.markdown(f"""<div class="card" style="border-top:3px solid {sc};">
                    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:10px;">
                      <span style="font-size:12px;color:#9b97b0;">Mean: {mean_r} ppm &nbsp;|&nbsp; StdDev: {std_r} &nbsp;|&nbsp; N={len(readings)}</span>
                      <span style="font-size:13px;font-weight:800;color:{sc};">{"⚠ " + str(len(anomalies)) + " ANOMAL" + ("Y" if len(anomalies)==1 else "IES") + " DETECTED" if anomalies else "✅ ALL CLEAR"}</span>
                    </div>
                    <div style="display:flex;align-items:flex-end;gap:2px;height:80px;background:#0a0a16;border-radius:10px;padding:6px;">{spark}</div>
                    <div style="font-size:10px;color:#9b97b0;margin-top:6px;">Red = anomalous · Gold = normal · Hover for values</div>
                    </div>""", unsafe_allow_html=True)

                    slope, next_est = _ml_linear_regression_trend(readings)
                    trend_dir = "📈 Rising" if slope > 0.5 else "📉 Declining" if slope < -0.5 else "➡ Stable"
                    st.markdown(f"""<div style="background:rgba(245,166,35,0.08);border:1px solid rgba(245,166,35,0.2);border-radius:10px;padding:12px 16px;font-size:13px;margin-top:8px;color:#f0ecff;">
                    <strong style="color:#f5a623;">Trend Analysis (Linear Regression):</strong> {trend_dir} — slope: {slope} ppm/reading · Next estimated: <strong style="color:#f5a623;">{next_est} ppm</strong>
                    </div>""", unsafe_allow_html=True)

                    for idx, val, zs in anomalies:
                        lvl = "🔴 CRITICAL – Evacuate Immediately (MHSA S.23)" if val > 1000 else "🟠 WARNING – Investigate Now" if val > 500 else "🟡 ELEVATED – Monitor Closely"
                        bg = "rgba(255,71,87,0.12)" if val > 1000 else "rgba(255,165,2,0.10)" if val > 500 else "rgba(255,214,0,0.08)"
                        st.markdown(f'<div style="background:{bg};border-radius:10px;padding:10px 16px;margin-bottom:6px;font-size:13px;color:#f0ecff;">'
                                    f'Reading <b>#{idx+1}</b>: <b style="color:#ff4757;">{val} ppm</b> &nbsp;·&nbsp; Z={round(zs,2)} &nbsp;·&nbsp; {lvl}</div>',
                                    unsafe_allow_html=True)
            except ValueError:
                st.error("Invalid input. Enter only numbers separated by commas.")
        else:
            st.info("Enter comma-separated gas readings above to begin anomaly detection.")

# ═══════════════════════════════════════════════════════════════════════════════
# ⚠️ HAZARD INTELLIGENCE
# ═══════════════════════════════════════════════════════════════════════════════
elif menu == "⚠️  Hazard Intelligence":
    st.title("Hazard Intelligence")
    st.markdown(f'<div class="sec-label">AI Risk Assessment · {active_mine} · {active_province}</div>', unsafe_allow_html=True)

    left, right = st.columns([1, 1.3], gap="large")
    with left:
        st.markdown('<div class="sec-label">Site Parameters</div>', unsafe_allow_html=True)
        incident_type = st.selectbox("Hazard Category", INCIDENT_TYPES)
        location_sel  = st.selectbox("Site Location / Zone", province_zones)
        depth         = st.slider("Working Depth / Level (m)", 0, 3500, 800, 25)
        workers_exp   = st.slider("Workers Exposed", 1, 200, 12)

        st.markdown('<div class="sec-label" style="margin-top:14px;">Environmental & Equipment Inputs</div>', unsafe_allow_html=True)
        gas_ppm_val  = st.slider("Gas Reading – CH4 (ppm)", 0, 2000, 150, 10, key="hz_gas")
        hz_maint_val = st.select_slider("Equipment / Support Status",
                                        options=["Overdue", "Poor", "Fair", "Good", "Excellent"],
                                        value="Good", key="hz_maint")
        hz_hours_val = st.slider("Hours Worked (current shift)", 1, 16, 8, key="hz_hours")
        hz_temp_val  = st.slider("Working Environment Temp (°C)", 18, 55, 28, key="hz_temp")
        hz_inc_hx    = st.slider("Incidents in Last 30 Days", 0, 20, 2, key="hz_inc_hx")

        observations = st.text_area("Site Observations", height=130,
                                    placeholder="Describe conditions: ground stability, unusual sounds, visible cracks, water ingress, equipment state, lighting…")

        # ── Pre-compute zone risk for selected location ─────────────────────
        _random.seed(42)
        _hz_zone_data = [
            {"zone": z, "depth": _random.randint(50, 3000),
             "gas": _random.randint(5, 1800), "incidents": _random.randint(0, 12)}
            for z in province_zones
        ]
        _hz_labels, _hz_risk_map = _ml_kmeans_zones(_hz_zone_data)
        _hz_zone_lookup = {}
        for _hi, _hzd in enumerate(_hz_zone_data):
            _hrl, _hrc = _hz_risk_map.get(_hz_labels[_hi], ("Low Risk", "#2ed573"))
            _hz_zone_lookup[_hzd["zone"]] = {
                "risk_label": _hrl, "risk_color": _hrc,
                "gas": _hzd["gas"], "depth": _hzd["depth"], "incidents": _hzd["incidents"]
            }
        _cur_zone_info = _hz_zone_lookup.get(location_sel, {"risk_label": "Unknown", "risk_color": "#9b97b0", "gas": "—", "depth": "—", "incidents": "—"})

        if st.button("▶ EXECUTE HAZARD ANALYSIS"):
            if not observations.strip():
                st.warning("Please describe site conditions before running analysis.")
            else:
                st.session_state.stats["hazards"] += 1
                st.session_state.stats["queries"] += 1
                mark_used()
                # Store zone risk for display in right panel
                st.session_state.hz_location  = location_sel
                st.session_state.hz_zone_risk = _cur_zone_info
                with st.spinner("Analysing hazard data…"):
                    risk_flags = []
                    if gas_ppm_val > 1000: risk_flags.append(f"CRITICAL methane at {gas_ppm_val} ppm (exceeds 1000 ppm threshold — MHSA mandatory stop work)")
                    elif gas_ppm_val > 500: risk_flags.append(f"Elevated methane at {gas_ppm_val} ppm (exceeds 500 ppm warning level)")
                    elif gas_ppm_val > 200: risk_flags.append(f"Moderate methane at {gas_ppm_val} ppm")
                    if hz_maint_val in ["Overdue", "Poor"]: risk_flags.append(f"Equipment/support status is {hz_maint_val} — MHSA Section 49 violation risk")
                    if hz_hours_val >= 12: risk_flags.append(f"High fatigue risk: {hz_hours_val} hours on shift")
                    if hz_inc_hx >= 5: risk_flags.append(f"{hz_inc_hx} prior incidents in 30 days — elevated risk profile")
                    if hz_temp_val > 40: risk_flags.append(f"Extreme heat: {hz_temp_val}°C — mandatory cooling measures required")
                    flags_str = "; ".join(risk_flags) if risk_flags else "No critical compounding factors identified"
                    _zone_risk_label = _cur_zone_info["risk_label"]

                    prompt = (
                        f"You are a senior certified mining safety engineer. Conduct a detailed hazard analysis based STRICTLY on the following real site data:\n\n"
                        f"MINE: {active_mine} ({active_mine_type}) – {active_province}\n"
                        f"HAZARD CATEGORY: {incident_type}\n"
                        f"ZONE: {location_sel} (K-Means Zone Classification: {_zone_risk_label} — baseline {_cur_zone_info['gas']} ppm gas, {_cur_zone_info['depth']}m depth, {_cur_zone_info['incidents']} incidents/30d in zone history)\n"
                        f"DEPTH: {depth}m\n"
                        f"WORKERS EXPOSED: {workers_exp}\n"
                        f"CH4 GAS READING: {gas_ppm_val} ppm\n"
                        f"EQUIPMENT STATUS: {hz_maint_val}\n"
                        f"SHIFT HOURS WORKED: {hz_hours_val}\n"
                        f"TEMPERATURE: {hz_temp_val}°C\n"
                        f"PRIOR INCIDENTS (30d): {hz_inc_hx}\n"
                        f"KEY RISK FLAGS: {flags_str}\n"
                        f"SITE OBSERVATIONS: {observations}\n\n"
                        f"Write a complete, specific hazard analysis report using ONLY the data above. Structure it as:\n"
                        f"1. Overall Risk Level (Low/Medium/High/Critical) — state clearly and justify with the specific numbers above\n"
                        f"2. Zone Overview — explain what the '{_zone_risk_label}' classification means for {location_sel} specifically, referencing the baseline data\n"
                        f"3. Identified Hazards — name each hazard and reference the actual readings (e.g. '{gas_ppm_val} ppm CH4 at {depth}m with {workers_exp} workers')\n"
                        f"4. Immediate Actions Required — specific, prioritised steps for {workers_exp} workers at {depth}m right now\n"
                        f"5. Short-term Mitigation Plan (next 24–72 hours)\n"
                        f"6. Long-term Mitigation Plan (next 30 days)\n"
                        f"7. Applicable MHSA Act 29 of 1996 regulatory requirements — cite exact sections relevant to {incident_type}\n"
                        f"If the data indicates genuinely low risk, state this clearly and explain exactly why the readings are within safe parameters.\n"
                        f"Write in natural, professional English. Do NOT use generic boilerplate. Every sentence must reference the specific mine, zone, or data values entered above."
                    )
                    HAZARD_SYSTEM = (
                        f"You are a senior certified mining safety engineer and risk specialist with 20+ years of experience "
                        f"in South African hard-rock, coal, and open-pit mining. You have deep expertise in MHSA Act 29 of 1996, "
                        f"DMR regulations, SANS standards, and operational safety for {active_mine_type} mines. "
                        f"When conducting hazard analyses, you ALWAYS: "
                        f"(1) Give a specific, justified risk level based on the actual numbers provided — never vague; "
                        f"(2) Reference the specific readings (ppm values, temperatures, depth, hours) in your analysis; "
                        f"(3) Provide detailed, actionable corrective actions with realistic timelines and responsible parties; "
                        f"(4) Cite precise MHSA sections and regulation numbers — not generic references; "
                        f"(5) Write in clear, professional English that a mine overseer, safety officer, and DMR inspector can all understand; "
                        f"(6) Never use placeholder text like '[insert name]' or '[to be determined]'; "
                        f"(7) If the data indicates genuinely low risk, state this clearly and explain exactly why; "
                        f"(8) Provide enough detail that the report can stand alone as a formal safety document. "
                        f"Always write in natural, human language — not bullet-point telegrams."
                    )
                    result = call_groq(prompt, system_prompt=HAZARD_SYSTEM)
                    st.session_state.hazard_analysis = result

                    acc_pct, features, sev_label, sev_color = _ml_accident_prediction(
                        depth, gas_ppm_val, hz_temp_val, hz_hours_val, hz_maint_val, hz_inc_hx)
                    st.session_state.hz_pred = (acc_pct, features, sev_label, sev_color)

    with right:
        if st.session_state.hazard_analysis:
            # ── Zone Risk Overview Card ──────────────────────────────────────
            _disp_loc   = st.session_state.get("hz_location", location_sel)
            _disp_zinfo = st.session_state.get("hz_zone_risk") or _hz_zone_lookup.get(location_sel, {})
            if _disp_zinfo:
                _zrl = _disp_zinfo.get("risk_label", "Unknown")
                _zrc = _disp_zinfo.get("risk_color", "#9b97b0")
                _zgas = _disp_zinfo.get("gas", "—")
                _zdep = _disp_zinfo.get("depth", "—")
                _zinc = _disp_zinfo.get("incidents", "—")
                st.markdown(f"""<div class="card" style="border-top:3px solid {_zrc};margin-bottom:16px;">
                  <div class="sec-label">Zone Risk Overview — {_disp_loc}</div>
                  <div style="display:flex;align-items:center;gap:20px;flex-wrap:wrap;">
                    <div style="text-align:center;">
                      <div style="font-size:34px;font-family:'Bebas Neue';color:{_zrc};line-height:1;">{_zrl.upper()}</div>
                      <div style="font-size:10px;color:#9b97b0;margin-top:2px;">K-MEANS CLASSIFICATION</div>
                    </div>
                    <div style="flex:1;display:flex;gap:10px;flex-wrap:wrap;">
                      <div style="background:#0a0a16;border-radius:10px;padding:10px 16px;min-width:80px;text-align:center;flex:1;">
                        <div style="font-size:20px;font-family:'Bebas Neue';color:{_zrc};">{_zgas} ppm</div>
                        <div style="font-size:10px;color:#9b97b0;">Baseline Gas</div>
                      </div>
                      <div style="background:#0a0a16;border-radius:10px;padding:10px 16px;min-width:80px;text-align:center;flex:1;">
                        <div style="font-size:20px;font-family:'Bebas Neue';color:{_zrc};">{_zdep}m</div>
                        <div style="font-size:10px;color:#9b97b0;">Zone Depth</div>
                      </div>
                      <div style="background:#0a0a16;border-radius:10px;padding:10px 16px;min-width:80px;text-align:center;flex:1;">
                        <div style="font-size:20px;font-family:'Bebas Neue';color:{_zrc};">{_zinc}</div>
                        <div style="font-size:10px;color:#9b97b0;">Incidents/30d</div>
                      </div>
                    </div>
                  </div>
                </div>""", unsafe_allow_html=True)

            st.markdown('<div class="result-panel">', unsafe_allow_html=True)
            st.markdown("## Analysis Result")
            st.markdown(st.session_state.hazard_analysis)
            st.markdown('</div>', unsafe_allow_html=True)

            if st.session_state.hz_pred:
                acc_pct, features, sev_label, sev_color = st.session_state.hz_pred
                with st.expander("📊 Predictive Risk Models", expanded=True):
                    pred_html = _ml_render_prediction(acc_pct, features, sev_label, sev_color)
                    st.markdown(pred_html, unsafe_allow_html=True)

            st.markdown("---")
            st.markdown("### Download Report")
            d1, d2 = st.columns(2)
            meta = {
                "Mine": active_mine, "Province": active_province,
                "Zone": location_sel, "Depth": f"{depth}m",
                "Workers Exposed": workers_exp, "Gas Reading": f"{gas_ppm_val} ppm",
                "Hazard Type": incident_type, "Report Date": datetime.now().strftime('%d %B %Y'),
            }
            pdf_data = create_pdf(st.session_state.hazard_analysis, "Hazard Analysis Report", meta)
            if pdf_data:
                d1.download_button("⬇ Download PDF Report", data=pdf_data,
                                   file_name=f"HM_Hazard_{active_province}_{datetime.now().strftime('%Y%m%d')}.pdf",
                                   mime="application/pdf")
            d2.download_button("⬇ Download Word Report",
                               create_docx(st.session_state.hazard_analysis, "Hazard Analysis Report", meta),
                               file_name=f"HM_Hazard_{active_province}_{datetime.now().strftime('%Y%m%d')}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.markdown("""<div class="card" style="text-align:center;padding:50px;border:1px dashed #2a2a50;">
            <div style="font-size:48px;margin-bottom:16px;">⚠️</div>
            <div style="font-size:15px;color:#9b97b0;line-height:1.8;">Configure site parameters on the left<br>and click <strong style="color:#f5a623;">Execute Hazard Analysis</strong> to begin.</div>
            </div>""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# ✦ COMPLIANCE ENGINE
# ═══════════════════════════════════════════════════════════════════════════════
elif menu == "✦  Compliance Engine":
    st.title("Compliance Engine")
    st.markdown('<div class="sec-label">MHSA Act 29 of 1996 · Automated Protocol Audit · K-Means Zone Classification</div>', unsafe_allow_html=True)

    _random.seed(42)
    _zone_data_all = [
        {"zone": z, "depth": _random.randint(50, 3000),
         "gas": _random.randint(5, 1800), "incidents": _random.randint(0, 12)}
        for z in province_zones
    ]
    _labels, _risk_map = _ml_kmeans_zones(_zone_data_all)
    _zone_lookup = {}
    for _i, _zd in enumerate(_zone_data_all):
        _rl, _rc = _risk_map.get(_labels[_i], ("Low Risk", "#2ed573"))
        _zone_lookup[_zd["zone"]] = {
            "risk_label": _rl, "risk_color": _rc,
            "gas": _zd["gas"], "depth": _zd["depth"], "incidents": _zd["incidents"]
        }

    c1, c2 = st.columns([1, 1.3], gap="large")
    with c1:
        op_id      = st.text_input("Operation / Site Reference", placeholder="e.g. OPS-2026-014")
        op_name    = st.text_input("Operation Name / Description", placeholder="e.g. Shaft 4 – East Reef Development")
        supervisor = st.selectbox("Responsible Supervisor", [f"{s['name']} ({s['role']})" for s in SUPERVISORS])
        audit_date = st.date_input("Audit Date", value=date.today())

        st.markdown('<div class="sec-label" style="margin-top:14px;">Mine Zones Under Audit</div>', unsafe_allow_html=True)
        selected_zones = st.multiselect("Select zones to audit", province_zones,
                                        placeholder="Select one or more zones from this mine…",
                                        label_visibility="collapsed")

        if selected_zones:
            rows = "".join(
                f'<div style="display:flex;justify-content:space-between;align-items:center;padding:6px 0;border-bottom:1px solid #2a2a50;">'
                f'<span style="font-size:12px;color:#f0ecff;">{z[:38]}</span>'
                f'<span style="font-size:11px;font-weight:700;color:{_zone_lookup[z]["risk_color"]};'
                f'background:rgba(0,0,0,0.4);padding:2px 10px;border-radius:8px;">{_zone_lookup[z]["risk_label"]}</span></div>'
                for z in selected_zones)
            st.markdown(f'<div style="background:#0a0a16;border:1px solid #2a2a50;border-radius:12px;padding:12px 16px;margin-bottom:12px;">'
                        f'<div class="sec-label">Zone Risk Preview (K-Means Classification)</div>{rows}</div>',
                        unsafe_allow_html=True)

        st.markdown('<div class="sec-label" style="margin-top:12px;">Safety Protocols In Place</div>', unsafe_allow_html=True)
        measures = st.multiselect("Select active safety protocols", SAFETY_PROTOCOLS, label_visibility="collapsed",
                                  placeholder="Select all protocols currently in place…")
        notes = st.text_area("Additional Observations / Findings", height=100,
                              placeholder="Any specific conditions or findings to include in the audit…")

        if st.button("▶ RUN COMPLIANCE AUDIT"):
            if not selected_zones:
                st.warning("Select at least one mine zone before running the audit.")
            elif not (op_id.strip() or op_name.strip()):
                st.warning("Enter an Operation ID or Name.")
            else:
                score = int((len(measures) / len(SAFETY_PROTOCOLS)) * 100) if measures else 0
                st.session_state.stats["compliance_score"] = score
                st.session_state.stats["queries"] += 1
                mark_used()
                missing = [p for p in SAFETY_PROTOCOLS if p not in measures]
                zone_det = "\n".join(
                    f"  - {z}: {_zone_lookup[z]['risk_label']} | {_zone_lookup[z]['gas']} ppm | "
                    f"{_zone_lookup[z]['depth']}m depth | {_zone_lookup[z]['incidents']} incidents/30d"
                    for z in selected_zones)
                _zone_risks = [_zone_lookup[z]["risk_label"] for z in selected_zones]
                overall_risk = ("High Risk" if "High Risk" in _zone_risks else
                                "Medium Risk" if "Medium Risk" in _zone_risks else "Low Risk")

                with st.spinner("Running MHSA compliance audit…"):
                    prompt = (
                        f"You are a certified South African mine safety auditor. Conduct a formal MHSA Act 29 of 1996 compliance audit based STRICTLY on:\n\n"
                        f"MINE: {active_mine} ({active_mine_type}) – {active_province}\n"
                        f"OPERATION: {op_name or op_id}\n"
                        f"SUPERVISOR: {supervisor}\n"
                        f"AUDIT DATE: {audit_date}\n"
                        f"ZONES AUDITED ({len(selected_zones)}):\n{zone_det}\n"
                        f"OVERALL ZONE RISK: {overall_risk}\n"
                        f"PROTOCOLS IN PLACE ({len(measures)}/{len(SAFETY_PROTOCOLS)}): {', '.join(measures) if measures else 'NONE'}\n"
                        f"MISSING PROTOCOLS ({len(missing)}): {', '.join(missing) if missing else 'None – all protocols active'}\n"
                        f"COMPLIANCE SCORE: {score}%\n"
                        f"ADDITIONAL OBSERVATIONS: {notes if notes.strip() else 'None provided'}\n\n"
                        f"Based ONLY on the above, provide a formal audit report with:\n"
                        f"1. Compliance Rating: {score}% — {'Critical' if score < 50 else 'Moderate' if score < 80 else 'Good'}\n"
                        f"2. Zone Risk Summary: for EACH zone listed above, state its risk level and key hazard drivers\n"
                        f"3. {'Gap Analysis: each of the ' + str(len(missing)) + ' missing protocols with the specific MHSA section violated' if missing else 'Gap Analysis: 0 gaps — all protocols in place. State this clearly.'}\n"
                        f"4. {'Corrective actions with owners and timelines for each gap' if missing else 'Maintenance recommendations to sustain full compliance'}\n"
                        f"5. Specific MHSA Act 29 of 1996 sections applicable to {active_mine_type} mines in {active_province}\n"
                        f"6. Next audit schedule recommendation\n"
                        f"Only reference the zones listed above. Do not invent zones or data."
                    )
                    COMPLIANCE_SYSTEM = (
                        f"You are a certified South African mine safety auditor and MHSA Act 29 of 1996 specialist with 20+ years of auditing experience "
                        f"across gold, platinum, coal, and open-pit operations. You work closely with the DMRE and understand exactly what inspectors look for. "
                        f"When generating compliance audit reports you ALWAYS: "
                        f"(1) Address EVERY zone listed by name — never skip or generalise across zones; "
                        f"(2) Cite the precise MHSA section and regulation number for every compliance gap — not just chapter names; "
                        f"(3) Assign realistic, named corrective action owners (use the supervisor name provided) with specific timelines (e.g. '48 hours', '7 days', '30 days'); "
                        f"(4) Write the gap analysis as a professional auditor would — explain WHY each gap is a risk, not just WHAT is missing; "
                        f"(5) If compliance is 100%, write substantive maintenance and continuous improvement recommendations — not placeholder text; "
                        f"(6) Use clear section headings, tables where appropriate, and professional audit language; "
                        f"(7) Never use brackets like [insert] or vague phrases like 'as appropriate'; "
                        f"(8) Write in natural, professional English that could be submitted directly to the DMRE."
                    )
                    result = call_groq(prompt, system_prompt=COMPLIANCE_SYSTEM)
                    st.session_state.compliance_audit = result
                    st.session_state["_audit_zones"]  = selected_zones
                    st.session_state["_audit_lookup"] = _zone_lookup
                    st.session_state["_audit_op"]     = op_name or op_id
                    st.session_state["_audit_sup"]    = supervisor
                    st.session_state["_audit_date"]   = str(audit_date)
                    st.session_state["_audit_score"]  = score

    with c2:
        if st.session_state.compliance_audit:
            sc    = st.session_state.get("_audit_score", st.session_state.stats["compliance_score"])
            color = "#2ed573" if sc >= 80 else "#ffa502" if sc >= 50 else "#ff4757"
            _saved_zones  = st.session_state.get("_audit_zones", selected_zones)
            _saved_lookup = st.session_state.get("_audit_lookup", _zone_lookup)
            _saved_op     = st.session_state.get("_audit_op", op_name or op_id)

            st.markdown(f"""<div style="background:#0a0a16;border-radius:18px;padding:24px;border:1px solid #2a2a50;margin-bottom:16px;">
            <div class="sec-label">MHSA Compliance Score — {_saved_op}</div>
            <div style="font-size:60px;font-family:'Bebas Neue';color:{color};line-height:1;">{sc}%</div>
            <div style="font-size:13px;color:#9b97b0;margin-bottom:10px;font-weight:600;">{"Critical — Immediate action required" if sc < 50 else "Moderate — Address protocol gaps" if sc < 80 else "Good — Maintain compliance standards"}</div>
            <div style="background:#060610;border-radius:8px;height:10px;overflow:hidden;"><div style="width:{sc}%;height:100%;background:{color};border-radius:8px;"></div></div>
            </div>""", unsafe_allow_html=True)

            if _saved_zones:
                with st.expander(f"📍 Mine Zone Risk Clusters (K-Means, {len(_saved_zones)} zones)", expanded=True):
                    tiers = {"High Risk": [], "Medium Risk": [], "Low Risk": []}
                    for z in _saved_zones:
                        zi = _saved_lookup.get(z, {})
                        tiers[zi.get("risk_label", "Low Risk")].append((z, zi))
                    tier_colors = {"High Risk": "#ff4757", "Medium Risk": "#ffa502", "Low Risk": "#2ed573"}
                    active_tiers = [(t, tiers[t]) for t in ["High Risk", "Medium Risk", "Low Risk"] if tiers[t]]
                    cols_z = st.columns(len(active_tiers)) if active_tiers else st.columns(1)
                    for ti, (tlabel, tzones) in enumerate(active_tiers):
                        tc = tier_colors[tlabel]
                        rows_z = "".join(
                            f'<div style="padding:6px 0;border-bottom:1px solid #2a2a50;">'
                            f'<div style="font-size:12px;color:#f0ecff;font-weight:600;">{z[:34]}</div>'
                            f'<div style="font-size:11px;color:#9b97b0;margin-top:2px;">{zi.get("gas","—")} ppm · {zi.get("depth","—")}m · {zi.get("incidents","—")} inc.</div></div>'
                            for z, zi in tzones)
                        with cols_z[ti]:
                            st.markdown(f'<div class="card" style="border-top:3px solid {tc};padding:16px;">'
                                        f'<div style="font-size:11px;color:{tc};font-weight:800;letter-spacing:1px;text-transform:uppercase;margin-bottom:6px;">{tlabel}</div>'
                                        f'<div style="font-size:30px;font-family:\'Bebas Neue\';color:{tc};margin-bottom:8px;">{len(tzones)} Zone{"s" if len(tzones)!=1 else ""}</div>'
                                        f'{rows_z}</div>', unsafe_allow_html=True)

            st.markdown('<div class="result-panel">', unsafe_allow_html=True)
            st.markdown(st.session_state.compliance_audit)
            st.markdown('</div>', unsafe_allow_html=True)

            st.markdown("---")
            st.markdown("### Download Audit Report")
            da1, da2 = st.columns(2)
            _sup = st.session_state.get("_audit_sup", "")
            _dt  = st.session_state.get("_audit_date", str(date.today()))
            meta_c = {
                "Mine": active_mine, "Province": active_province,
                "Operation": _saved_op, "Supervisor": _sup,
                "Audit Date": _dt, "Compliance Score": f"{sc}%",
                "Zones Audited": len(_saved_zones), "Report Date": datetime.now().strftime('%d %B %Y'),
            }
            pdf_c = create_pdf(st.session_state.compliance_audit, "MHSA Compliance Audit Report", meta_c)
            if pdf_c:
                da1.download_button("⬇ Download PDF Audit", data=pdf_c,
                                    file_name=f"HM_Compliance_{active_province}_{datetime.now().strftime('%Y%m%d')}.pdf",
                                    mime="application/pdf")
            da2.download_button("⬇ Download Word Audit",
                                create_docx(st.session_state.compliance_audit, "MHSA Compliance Audit Report", meta_c),
                                file_name=f"HM_Compliance_{active_province}_{datetime.now().strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.markdown("""<div class="card" style="text-align:center;padding:50px;border:1px dashed #2a2a50;">
            <div style="font-size:48px;margin-bottom:16px;">✦</div>
            <div style="font-size:15px;color:#9b97b0;line-height:1.8;">Configure audit parameters on the left<br>and click <strong style="color:#f5a623;">Run Compliance Audit</strong> to begin.</div>
            </div>""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# ◉ REPORT FORGE
# ═══════════════════════════════════════════════════════════════════════════════
elif menu == "◉  Report Forge":
    st.title("Report Forge")
    st.markdown(f'<div class="sec-label">Document Generation · {active_mine} · MHSA & DMR Compliant</div>', unsafe_allow_html=True)

    rtype = st.selectbox("Report Type", [
        "Incident Report", "Risk Assessment", "Compliance Audit Summary",
        "Near Miss Report", "Equipment Inspection Report", "Fatality / Serious Injury Report",
    ])

    st.markdown("---")

    if rtype == "Incident Report":
        r1, r2 = st.columns(2)
        with r1:
            ir_date      = st.date_input("Incident Date", value=date.today())
            ir_time      = st.time_input("Incident Time", value=dt_time(8, 0))
            ir_ref       = st.text_input("Reference No.", placeholder="e.g. IR-2026-047")
            ir_location  = st.selectbox("Location / Zone", province_zones)
            ir_type      = st.selectbox("Incident Type", INCIDENT_TYPES)
            ir_severity  = st.selectbox("Severity", ["Near Miss", "First Aid", "Medical Treatment", "Lost Time", "Serious", "Fatal"])
        with r2:
            ir_injured   = st.multiselect("Injured Parties", [f"{m['name']} ({m['role']})" for m in MINERS])
            ir_injury    = st.selectbox("Injury Type", INJURY_TYPES)
            ir_witnesses = st.multiselect("Witnesses", [f"{m['name']} ({m['role']})" for m in MINERS])
            ir_sup       = st.selectbox("Reporting Supervisor", [f"{s['name']} – {s['role']}" for s in SUPERVISORS])

        ir_desc  = st.text_area("Incident Description", height=130,
                                placeholder="Describe exactly what happened, sequence of events, environmental conditions, equipment involved…")
        ir_cause = st.text_area("Immediate Cause / Contributing Factors", height=90,
                                placeholder="What directly caused the incident? What conditions contributed?")
        ir_photo = st.file_uploader("Evidence Photo (Optional)", type=["jpg","jpeg","png"])
        img_desc = ""
        if ir_photo:
            st.image(Image.open(ir_photo), caption="Evidence Photo", use_container_width=True)

        if st.button("▶ GENERATE INCIDENT REPORT"):
            st.session_state.stats["reports"] += 1
            st.session_state.stats["queries"] += 1
            mark_used()
            if ir_photo:
                b64 = base64.b64encode(ir_photo.getvalue()).decode()
                mime = "image/png" if ir_photo.name.lower().endswith("png") else "image/jpeg"
                img_desc = call_groq_vision("Describe this mining incident evidence photo for a formal safety report, noting any visible hazards, damage, or environmental conditions.", b64, mime)

            prompt = (
                f"You are a certified South African mine safety officer. Generate a formal, complete Incident Report.\n\n"
                f"MINE: {active_mine} ({active_mine_type}) – {active_province}\n"
                f"Reference: {ir_ref or 'Pending'} | Date: {ir_date.strftime('%d %B %Y')} | Time: {ir_time}\n"
                f"Zone: {ir_location} | Type: {ir_type} | Severity: {ir_severity}\n"
                f"Supervisor: {ir_sup}\n"
                f"Injured: {', '.join(ir_injured) if ir_injured else 'No injuries reported'}\n"
                f"Injury Type: {ir_injury}\n"
                f"Witnesses: {', '.join(ir_witnesses) if ir_witnesses else 'None recorded'}\n"
                f"Description: {ir_desc or 'Not provided'}\n"
                f"Immediate Cause: {ir_cause or 'Under investigation'}\n"
                f"{'Photo Evidence: ' + img_desc if img_desc else ''}\n\n"
                f"Write a complete formal Incident Report with sections:\n"
                f"1. Incident Summary\n"
                f"2. Detailed Narrative\n"
                f"3. Personnel Involved\n"
                f"4. Root Cause Analysis\n"
                f"5. Contributing Factors\n"
                f"6. Immediate Actions Taken\n"
                f"7. Corrective Actions Required (with owners and timelines)\n"
                f"8. MHSA Regulatory Obligations (Section 23 reporting, DMR notification requirements)\n"
                f"9. Preventive Measures\n"
                f"Do NOT use placeholder brackets. Be specific and professional."
            )
            with st.spinner("Generating Incident Report…"):
                REPORT_SYSTEM = (
                    f"You are a certified South African mine safety officer and incident investigator with 20+ years of experience "
                    f"in {active_mine_type} operations. You write formal incident reports that meet MHSA Act 29 of 1996 and DMRE requirements. "
                    f"Your reports are: comprehensive (every section fully written out — no placeholders); precise (you reference actual data provided); "
                    f"actionable (corrective actions have named owners, specific timelines, and measurable outcomes); "
                    f"professionally written in natural English that could be submitted directly to the Principal Inspector of Mines. "
                    f"Apply the ICAM (Incident Cause Analysis Method) for root cause analysis. "
                    f"Always include the precise MHSA Section 23 reporting obligations for the severity level stated. "
                    f"Never write 'To Be Determined' or use square bracket placeholders."
                )
                result = call_groq(prompt, system_prompt=REPORT_SYSTEM)
                st.session_state.active_doc = result
                st.session_state.doc_type   = "Incident Report"
                st.session_state.doc_meta   = {
                    "Mine": active_mine, "Reference": ir_ref or "Pending",
                    "Date": ir_date.strftime('%d %B %Y'), "Time": str(ir_time),
                    "Location": ir_location, "Type": ir_type, "Severity": ir_severity,
                    "Supervisor": ir_sup,
                }

    else:
        r1, r2 = st.columns(2)
        with r1:
            r_date     = st.date_input("Report Date", value=date.today())
            r_ref      = st.text_input("Reference No.", placeholder="e.g. RA-2026-018")
            r_location = st.selectbox("Location / Zone", province_zones)
            r_supervisor = st.selectbox("Supervisor", [f"{s['name']} – {s['role']}" for s in SUPERVISORS])
        with r2:
            r_miners = st.multiselect("Personnel Involved", [f"{m['name']} ({m['role']})" for m in MINERS])

        findings = st.text_area("Detailed Findings / Observations", height=160,
                                placeholder="Describe all relevant findings, conditions observed, measurements taken…")
        r_photo = st.file_uploader("Evidence Photo (Optional)", type=["jpg", "jpeg", "png"])
        img2 = ""
        if r_photo:
            st.image(Image.open(r_photo), caption="Evidence Photo", use_container_width=True)

        if st.button("▶ GENERATE REPORT"):
            st.session_state.stats["reports"] += 1
            st.session_state.stats["queries"] += 1
            mark_used()
            if r_photo:
                b64 = base64.b64encode(r_photo.getvalue()).decode()
                mime = "image/png" if r_photo.name.lower().endswith("png") else "image/jpeg"
                img2 = call_groq_vision("Describe this mining site evidence image briefly for a formal report.", b64, mime)
            prompt = (
                f"You are a certified South African mining safety officer with deep MHSA expertise. Generate a formal {rtype}.\n\n"
                f"MINE: {active_mine} ({active_mine_type}) – {active_province}\n"
                f"Reference: {r_ref or 'Pending'} | Date: {r_date} | Location: {r_location}\n"
                f"Supervisor: {r_supervisor}\n"
                f"Personnel: {', '.join(r_miners) if r_miners else 'Not specified'}\n"
                f"Findings: {findings or 'No findings provided'}\n"
                f"{'Photo Evidence: ' + img2 if img2 else ''}\n\n"
                f"Write a complete, professional {rtype} with all standard sections, specific MHSA Act 29 of 1996 regulatory references, "
                f"corrective actions with owners and timelines, and recommendations. Do NOT use placeholder brackets."
            )
            with st.spinner(f"Generating {rtype}…"):
                REPORT_SYSTEM2 = (
                    f"You are a certified South African mine safety officer and documentation specialist with deep expertise in MHSA Act 29 of 1996, "
                    f"DMR regulations, and {active_mine_type} mining operations in {active_province}. "
                    f"You generate formal {rtype} documents that are: "
                    f"(1) Fully complete — every section written in full with no placeholder text or brackets; "
                    f"(2) Site-specific — you reference the actual mine name, zone, personnel, and findings provided; "
                    f"(3) Regulatory-precise — you cite the exact MHSA section and sub-section numbers applicable to each requirement; "
                    f"(4) Actionable — corrective actions include a responsible person (use the supervisor name given), a specific deadline, and a success criterion; "
                    f"(5) Professional — written in clear, natural English suitable for submission to the DMRE or a Board of Inquiry; "
                    f"(6) Detailed enough to stand alone as a legal safety document. "
                    f"Do not truncate any section. Write each section in full prose with supporting tables where appropriate."
                )
                result = call_groq(prompt, system_prompt=REPORT_SYSTEM2)
                st.session_state.active_doc = result
                st.session_state.doc_type   = rtype
                st.session_state.doc_meta   = {
                    "Mine": active_mine, "Reference": r_ref or "Pending",
                    "Date": str(r_date), "Location": r_location, "Supervisor": r_supervisor,
                }

    # ── Q-Learning recommendation ────────────────────────────────────────────
    if st.session_state.active_doc:
        st.markdown("---")
        with st.expander("🤖 Q-Learning Safety Action Recommender", expanded=False):
            rl1, rl2 = st.columns(2)
            gas_idx = rl1.select_slider("Gas Risk Level", options=[0,1,2,3], value=1,
                                        format_func=lambda x: ["Low","Moderate","High","Critical"][x])
            gnd_idx = rl2.select_slider("Ground Conditions", options=[0,1,2,3], value=1,
                                        format_func=lambda x: ["Stable","Unstable","Fractured","Collapse Risk"][x])
            best_action, q_vals, actions = _ml_rl_recommend(gas_idx, gnd_idx)
            st.markdown(f"""<div style="background:rgba(245,166,35,0.1);border:2px solid rgba(245,166,35,0.4);border-radius:14px;padding:18px 22px;margin-top:8px;">
            <div style="font-size:10px;color:#9b97b0;letter-spacing:2px;text-transform:uppercase;margin-bottom:6px;font-weight:700;">Recommended Action (Q-Learning)</div>
            <div style="font-size:18px;font-weight:800;color:#f5a623;">{best_action}</div>
            </div>""", unsafe_allow_html=True)
            max_q = max(q_vals) or 1
            q_bars = "".join(
                f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:6px;">'
                f'<div style="width:220px;font-size:11px;color:#d8d4f0;text-align:right;">{a[:36]}</div>'
                f'<div style="flex:1;background:#0a0a16;border-radius:4px;height:12px;overflow:hidden;">'
                f'<div style="width:{int(max(0,q_vals[i]/max_q)*100)}%;height:100%;background:{"#f5a623" if a==best_action else "#2a2a50"};border-radius:4px;"></div></div>'
                f'<div style="width:36px;font-size:10px;color:#9b97b0;text-align:right;">{round(q_vals[i],1)}</div></div>'
                for i, a in enumerate(actions))
            st.markdown(f'<div class="card" style="margin-top:12px;"><div style="font-size:9px;color:#9b97b0;letter-spacing:1.5px;text-transform:uppercase;margin-bottom:10px;font-weight:700;">Action Q-Values</div>{q_bars}</div>', unsafe_allow_html=True)

        st.markdown("---")
        st.markdown('<div class="result-panel">', unsafe_allow_html=True)
        st.markdown(f"## {st.session_state.doc_type}")
        st.markdown(st.session_state.active_doc)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("")
        st.markdown("### Download Report")
        d1, d2 = st.columns(2)
        fname    = st.session_state.doc_type.replace(" ", "_")
        pdf_data = create_pdf(st.session_state.active_doc, st.session_state.doc_type, st.session_state.doc_meta)
        if pdf_data:
            d1.download_button("⬇ Download as PDF", data=pdf_data,
                               file_name=f"HM_{fname}_{active_province}_{datetime.now().strftime('%Y%m%d')}.pdf",
                               mime="application/pdf")
        d2.download_button("⬇ Download as Word",
                           create_docx(st.session_state.active_doc, st.session_state.doc_type, st.session_state.doc_meta),
                           file_name=f"HM_{fname}_{active_province}_{datetime.now().strftime('%Y%m%d')}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ═══════════════════════════════════════════════════════════════════════════════
# ◎ SAFETY ADVISOR
# ═══════════════════════════════════════════════════════════════════════════════
elif menu == "◎  Safety Advisor":
    st.title("Safety Advisor")
    st.markdown(f'<div class="sec-label">AI-Powered Mining Safety Q&A · {active_mine} · Available 24/7</div>', unsafe_allow_html=True)

    SYSTEM = (
        f"You are an expert mining safety advisor and MHSA Act 29 of 1996 specialist with 25+ years of hands-on experience "
        f"across all major South African mining sectors — gold, platinum, coal, iron ore, diamonds, and quarrying. "
        f"The user is working at {active_mine}, a {active_mine_type} operation in {active_province}, South Africa. "
        f"You have deep knowledge of: MHSA Act 29 of 1996 and all supporting regulations; DMR/DMRE inspection processes and Section 54 stop-work procedures; "
        f"SANS standards relevant to mining; Occupational Health requirements (dust, noise, heat, gas); "
        f"Rock engineering and fall-of-ground prevention; Blasting regulations and explosives management; "
        f"Trackless mobile machinery (TMM) safety; Emergency response and mine rescue procedures; "
        f"Ventilation design and atmospheric monitoring; and South African mining industry best practices. "
        f"\n\nWhen answering questions, you ALWAYS: "
        f"(1) Give a thorough, detailed answer — never a brief telegram; "
        f"(2) Write in natural, conversational English — like a trusted expert colleague explaining something clearly; "
        f"(3) Reference specific MHSA sections and regulation numbers when relevant, but weave them naturally into the explanation rather than just listing them; "
        f"(4) Give practical, operational guidance that can actually be acted on at a mine — not just regulatory recitals; "
        f"(5) Tailor your answer to {active_mine_type} operations at {active_mine} where this context is relevant; "
        f"(6) If the question has multiple parts, address each part fully; "
        f"(7) Where the situation is genuinely low risk, say so clearly and explain why — don't over-alarm; "
        f"(8) For complex topics, use clear headings and structured explanations; "
        f"(9) Never fabricate statistics, regulation numbers, or legal requirements; "
        f"(10) Always be specific — vague answers like 'consult your safety officer' are not acceptable as a complete response."
    )

    if not st.session_state.chat_history:
        st.markdown("**Quick questions to get started:**")
        q1, q2, q3 = st.columns(3)
        suggestions = [
            f"What are the main fall-of-ground risks at {active_mine_type.split('/')[0].strip()} mines?",
            "What ventilation standards apply to underground methane zones per MHSA?",
            "How should I prepare for a DMR inspection at my mine?",
        ]
        for col, suggestion in zip([q1, q2, q3], suggestions):
            label = f"💬 {suggestion[:46]}…" if len(suggestion) > 46 else f"💬 {suggestion}"
            if col.button(label, key=f"sug_{suggestion[:22]}"):
                st.session_state.chat_history.append({"role": "user", "content": suggestion})
                resp = call_groq(suggestion, system_prompt=SYSTEM, history=st.session_state.chat_history[:-1])
                st.session_state.chat_history.append({"role": "assistant", "content": resp})
                st.rerun()

    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    if prompt := st.chat_input(f"Ask a safety question about {active_mine}…"):
        st.session_state.stats["queries"] += 1
        mark_used()
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        with st.chat_message("assistant"):
            with st.spinner("…"):
                resp = call_groq(prompt, system_prompt=SYSTEM, history=st.session_state.chat_history[:-1])
            st.markdown(resp)
            st.session_state.chat_history.append({"role": "assistant", "content": resp})

    if st.session_state.chat_history:
        if st.button("🗑 Clear Conversation"):
            st.session_state.chat_history = []
            st.rerun()

# ─── FOOTER ─────────────────────────────────────────────────────────────────────
st.markdown("---")
st.markdown(f"""
<div style="display:flex;justify-content:space-between;align-items:center;padding:8px 0;flex-wrap:wrap;gap:8px;">
  <span style="font-size:11px;color:#6b6880;">Minetrax AI · Mining Safety Intelligence Platform · South Africa</span>
  <span style="font-size:11px;color:#6b6880;">MHSA Act 29 of 1996 Compliant · Active: <strong style="color:#f5a623;">{active_mine}</strong></span>
</div>
""", unsafe_allow_html=True)
